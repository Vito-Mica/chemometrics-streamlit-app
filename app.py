
import io, difflib, warnings, os
from pathlib import Path
warnings.filterwarnings("ignore")
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.cm as mcm
import matplotlib.colors as mcolors
import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from scipy import stats as sp_stats
from scipy.cluster.hierarchy import dendrogram, linkage
from sklearn.decomposition import PCA
from sklearn.cross_decomposition import PLSRegression
from sklearn.impute import SimpleImputer
from sklearn.preprocessing import StandardScaler, label_binarize
from sklearn.pipeline import Pipeline
from sklearn.base import clone
from sklearn.model_selection import (train_test_split, StratifiedKFold,
                                      cross_validate, permutation_test_score)
from sklearn.metrics import (accuracy_score, balanced_accuracy_score,
                              f1_score, confusion_matrix, classification_report,
                              silhouette_score)
from sklearn.discriminant_analysis import LinearDiscriminantAnalysis as LDA
from sklearn.neighbors import KNeighborsClassifier
from sklearn.svm import SVC
from sklearn.ensemble import RandomForestClassifier
from sklearn.cluster import KMeans, AgglomerativeClustering
try:
    import umap as umap_lib
    UMAP_AVAILABLE = True
except ImportError:
    UMAP_AVAILABLE = False

st.set_page_config(page_title="ChemoTrace · Multivariate Provenance Analysis", page_icon="🔬", layout="wide")

# ─── PASSWORD GATE ─────────────────────────────────────────────────────────────
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False

if not st.session_state.authenticated:
    _login_col = st.columns([1, 2, 1])[1]
    with _login_col:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(
            "<h2 style='text-align:center; color:#2c3e6b;'>🔬 ChemoTrace</h2>"
            "<p style='text-align:center; color:#6c7a89;'>Multivariate Provenance Analysis — Please enter password.</p>",
            unsafe_allow_html=True,
        )
        with st.form("login_form"):
            _pw = st.text_input("Password", type="password", placeholder="Enter password …")
            _submit = st.form_submit_button("🔓 Login", use_container_width=True)
        if _submit:
            try:
                if _pw == st.secrets["APP_PASSWORD"]:
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("❌ Wrong password.")
            except KeyError:
                st.error("⚠️ `APP_PASSWORD` is not configured in Streamlit Secrets. "
                         "Please create `.streamlit/secrets.toml`.")
    st.stop()

# ─── DEMO FILE PATH ───────────────────────────────────────────────────────────
DEMO_FILE = Path(__file__).resolve().parent / "data" / "demo.xlsx"

# ─── PREMIUM SCIENTIFIC THEME ─────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Base palette: dark blue, graphite, silver, scientific grey ── */
:root {
    --mineral-dark: #1a1f2e;
    --mineral-blue: #2c3e6b;
    --mineral-accent: #4a6fa5;
    --mineral-silver: #b8c4d0;
    --mineral-light: #e8ecf1;
    --mineral-gold: #c9a959;
    --mineral-surface: #f4f6f9;
}

/* Sidebar styling */
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1a1f2e 0%, #2c3e6b 100%);
}
section[data-testid="stSidebar"] .stMarkdown,
section[data-testid="stSidebar"] label,
section[data-testid="stSidebar"] .stRadio label,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] p {
    color: #dce3ec !important;
}
section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: #e8ecf1 !important;
    letter-spacing: 0.03em;
}
section[data-testid="stSidebar"] hr {
    border-color: rgba(184, 196, 208, 0.2);
}

/* Main content area refinements */
.main .block-container {
    padding-top: 2rem;
    max-width: 1200px;
}

/* Headings */
.main h1 {
    color: #1a1f2e;
    font-weight: 700;
    letter-spacing: -0.01em;
    border-bottom: 3px solid #4a6fa5;
    padding-bottom: 0.4rem;
    margin-bottom: 1rem;
}
.main h2 {
    color: #2c3e6b;
    font-weight: 600;
}
.main h3 {
    color: #3a506b;
    font-weight: 600;
}

/* Metric cards */
[data-testid="stMetric"] {
    background: linear-gradient(135deg, #f4f6f9 0%, #e8ecf1 100%);
    border: 1px solid #b8c4d0;
    border-radius: 8px;
    padding: 12px 16px;
    box-shadow: 0 1px 4px rgba(26, 31, 46, 0.06);
}
[data-testid="stMetricLabel"] {
    color: #3a506b !important;
    font-weight: 600;
    text-transform: uppercase;
    font-size: 0.72rem !important;
    letter-spacing: 0.06em;
}
[data-testid="stMetricValue"] {
    color: #1a1f2e !important;
    font-weight: 700;
}

/* Buttons */
.stButton > button[kind="primary"],
.stButton > button[data-testid="stBaseButton-primary"] {
    background: linear-gradient(135deg, #2c3e6b 0%, #4a6fa5 100%);
    color: white;
    border: none;
    font-weight: 600;
    letter-spacing: 0.02em;
    transition: all 0.2s ease;
    border-radius: 6px;
}
.stButton > button[kind="primary"]:hover,
.stButton > button[data-testid="stBaseButton-primary"]:hover {
    background: linear-gradient(135deg, #1a1f2e 0%, #2c3e6b 100%);
    box-shadow: 0 4px 12px rgba(44, 62, 107, 0.3);
}

/* Dataframes */
[data-testid="stDataFrame"] {
    border: 1px solid #dce3ec;
    border-radius: 6px;
}

/* Info/warning/success/error boxes */
.stAlert {
    border-radius: 6px;
}

/* Tab styling */
.stTabs [data-baseweb="tab-list"] {
    gap: 4px;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 6px 6px 0 0;
    font-weight: 500;
}

/* Download buttons */
.stDownloadButton > button {
    border: 1px solid #b8c4d0;
    border-radius: 6px;
    font-size: 0.85rem;
    transition: all 0.15s ease;
}
.stDownloadButton > button:hover {
    border-color: #4a6fa5;
    color: #2c3e6b;
}

/* Expanders */
.streamlit-expanderHeader {
    font-weight: 600;
    color: #2c3e6b;
}

/* Selectbox and slider refinements */
.stSelectbox label, .stSlider label, .stMultiSelect label,
.stCheckbox label, .stRadio label, .stNumberInput label {
    font-weight: 500;
    color: #3a506b;
}

/* Subtle divider */
hr {
    border-color: #dce3ec;
}
</style>
""", unsafe_allow_html=True)

# ─── CONSTANTS ────────────────────────────────────────────────────────────────
HERKUNFT_COL = "Herkunftsort"
CLASS_LEVELS = ["Mineralklasse", "Land", "Nummer"]
_BASE_COLORS = {
    "Muscovite": "#1a6faf", "Phlogopite": "#d4720b",
    "Zulieferer": "#c0392b", "Madagaskar": "#1e8a3c",
    "Indien": "#1a6faf", "China": "#7b2fa8", "default": "#7f7f7f",
}
ALL_MARKERS   = ["o","s","^","D","v","P","X","*","<",">","h","p","H","8","d"]
PLOTLY_SYMS   = ["circle","square","diamond","triangle-up","triangle-down",
                 "pentagon","hexagon","star","cross","x","circle-open",
                 "square-open","diamond-open","triangle-up-open","star-open"]
TAB10 = list(plt.cm.tab10.colors)

# ─── SESSION STATE ─────────────────────────────────────────────────────────────
_INIT = {
    "df_raw": None, "df_klassen": None, "df_merged": None,
    "df_num": None, "df_proc": None,
    "feature_cols": None, "lod_map": None,
    "class_level": "Mineralklasse", "symbol_detail": True,
    "pca_scores": None, "pca_loadings": None,
    "df_ratios": None, "df_grouped": None,
    "pca_group_scores": None, "pca_group_loadings": None,
    "mz_df_unknown": None, "mz_unk_name_col": None,
    "mz_state": None,
    "data_source": "upload",          # "upload" or "demo"
    "demo_loaded": False,             # tracks if demo was auto-loaded
}
for _k, _v in _INIT.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

# ─── FUZZY JOIN ───────────────────────────────────────────────────────────────
def fuzzy_match(name, choices, cutoff=0.55):
    nl = str(name).strip().lower()
    cl = [str(c).strip().lower() for c in choices]
    m = difflib.get_close_matches(nl, cl, n=1, cutoff=cutoff)
    return choices[cl.index(m[0])] if m else None

def join_klassen(df_daten, df_klassen):
    kl_names = df_klassen[HERKUNFT_COL].tolist()
    kl_map = {}; match_report = []
    for name in df_daten[HERKUNFT_COL]:
        match = fuzzy_match(name, kl_names)
        if match:
            row = df_klassen[df_klassen[HERKUNFT_COL] == match].iloc[0]
            kl_map[name] = {c: str(row[c]) for c in CLASS_LEVELS}
            match_report.append({"Daten": name, "Klassen": match,
                                  "Match": "✅ exakt" if name==match else "🔄 fuzzy"})
        else:
            kl_map[name] = {c: "Unbekannt" for c in CLASS_LEVELS}
            match_report.append({"Daten": name, "Klassen": "—", "Match": "❌ kein Match"})
    df = df_daten.copy()
    for col in CLASS_LEVELS:
        df[col] = df[HERKUNFT_COL].map(lambda n, c=col: kl_map.get(n, {}).get(c, "?"))
    return df, pd.DataFrame(match_report)

# ─── PALETTE ──────────────────────────────────────────────────────────────────
def get_palette(class_level, df, symbol_detail):
    if df is None or HERKUNFT_COL not in df.columns:
        return {}
    unique_cls = sorted(df[class_level].astype(str).unique())
    class_base = {}
    if class_level == "Nummer":
        cmap = mcm.get_cmap("tab10", max(len(unique_cls), 1))
        for i, cls in enumerate(unique_cls):
            class_base[cls] = mcolors.to_hex(cmap(i / max(len(unique_cls)-1, 1)))
    else:
        for cls in unique_cls:
            class_base[cls] = _BASE_COLORS.get(cls, _BASE_COLORS["default"])
    palette = {}
    for cls in unique_cls:
        members = sorted(df[df[class_level].astype(str) == cls][HERKUNFT_COL].unique())
        n = len(members); base = np.array(mcolors.to_rgb(class_base[cls]))
        for j, h in enumerate(members):
            t = 0.0 if n == 1 else (j/(n-1))*0.40
            col = mcolors.to_hex((1-t)*base + t*np.ones(3))
            mkr = ALL_MARKERS[j % len(ALL_MARKERS)] if symbol_detail else "o"
            palette[h] = (col, mkr)
    return palette

def pal_color_map(pal):
    return {h: c for h, (c, _) in pal.items()}

def pal_symbol_map(pal):
    return {h: PLOTLY_SYMS[ALL_MARKERS.index(m) % len(PLOTLY_SYMS)] for h, (_, m) in pal.items()}

def cur_pal():
    return get_palette(st.session_state.class_level,
                       st.session_state.df_merged,
                       st.session_state.symbol_detail)

# ─── FIG HELPERS ──────────────────────────────────────────────────────────────
def fig2bytes(fig, fmt):
    buf = io.BytesIO(); fig.savefig(buf, format=fmt, dpi=300, bbox_inches="tight")
    buf.seek(0); return buf.read()

def pub_dl(fig, stem):
    st.markdown("**⬇️ Publication Format:**")
    c1,c2,c3,_ = st.columns([1,1,1,5])
    with c1: st.download_button("📄 PDF", fig2bytes(fig,"pdf"), f"{stem}.pdf",
                                 "application/pdf", key=f"pdf_{stem}")
    with c2: st.download_button("🖼️ PNG", fig2bytes(fig,"png"), f"{stem}.png",
                                 "image/png", key=f"png_{stem}")
    with c3: st.download_button("✏️ SVG", fig2bytes(fig,"svg"), f"{stem}.svg",
                                 "image/svg+xml", key=f"svg_{stem}")

def scatter_mpl(df, xcol, ycol, pal, cl, title, fs=11, figsize=(8,6)):
    fig, ax = plt.subplots(figsize=figsize); done = set()
    for _, row in df.iterrows():
        h = str(row[HERKUNFT_COL]); col, mkr = pal.get(h, ("#aaa","o"))
        lbl = h if h not in done else "_nolegend_"; done.add(h)
        ax.scatter(row[xcol], row[ycol], color=col, marker=mkr, s=60,
                   alpha=0.88, label=lbl, lw=0.4, edgecolors="white", zorder=3)
    ax.set_xlabel(xcol, fontsize=fs); ax.set_ylabel(ycol, fontsize=fs)
    ax.set_title(title, fontsize=fs+1, fontweight="bold")
    ax.grid(True, alpha=0.18, lw=0.6)
    ax.legend(bbox_to_anchor=(1.02,1), loc="upper left", fontsize=max(fs-3,7), framealpha=0.85)
    fig.tight_layout(); return fig

def pca_mpl(Z, df_ref, pal, cl, title, ve=None, extras=None, fs=11, figsize=(8,6)):
    fig, ax = plt.subplots(figsize=figsize); done = set()
    for i, (_, row) in enumerate(df_ref.iterrows()):
        h = str(row[HERKUNFT_COL]); col, mkr = pal.get(h, ("#aaa","o"))
        lbl = h if h not in done else "_nolegend_"; done.add(h)
        ax.scatter(Z[i,0], Z[i,1], color=col, marker=mkr, s=60,
                   alpha=0.88, label=lbl, lw=0.4, edgecolors="white", zorder=3)
    if extras:
        for lbl, (Ze, col, mkr, sz) in extras.items():
            ax.scatter(Ze[:,0], Ze[:,1], color=col, marker=mkr, s=sz,
                       alpha=0.92, label=lbl, edgecolors="black", lw=0.8, zorder=5)
    xl = f"PC1 ({ve[0]*100:.1f}%)" if ve is not None else "PC1"
    yl = f"PC2 ({ve[1]*100:.1f}%)" if ve is not None else "PC2"
    ax.set_xlabel(xl, fontsize=fs); ax.set_ylabel(yl, fontsize=fs)
    ax.set_title(title, fontsize=fs+1, fontweight="bold")
    ax.grid(True, alpha=0.18, lw=0.6)
    ax.legend(bbox_to_anchor=(1.02,1), loc="upper left", fontsize=max(fs-3,7), framealpha=0.85)
    fig.tight_layout(); return fig

def biplot_mpl(Z, ld_df, df_ref, pal, cl, title, ve, top_n=8, scale=3.0, fs=11, figsize=(9,7)):
    fig, ax = plt.subplots(figsize=figsize); done = set()
    for i, (_, row) in enumerate(df_ref.iterrows()):
        h = str(row[HERKUNFT_COL]); col, mkr = pal.get(h, ("#aaa","o"))
        lbl = h if h not in done else "_nolegend_"; done.add(h)
        ax.scatter(Z[i,0], Z[i,1], color=col, marker=mkr, s=55,
                   alpha=0.85, label=lbl, lw=0.4, edgecolors="white", zorder=3)
    top_f = ld_df["PC1"].abs().add(ld_df["PC2"].abs()).nlargest(top_n).index
    for feat in top_f:
        lx = ld_df.loc[feat,"PC1"]*scale; ly = ld_df.loc[feat,"PC2"]*scale
        ax.annotate("", xy=(lx,ly), xytext=(0,0),
                    arrowprops=dict(arrowstyle="->", color="#555", lw=1.5))
        ax.text(lx*1.08, ly*1.08, feat, fontsize=8, color="#333", ha="center", va="center")
    ax.axhline(0, color="grey", lw=0.5, ls="--"); ax.axvline(0, color="grey", lw=0.5, ls="--")
    ax.set_xlabel(f"PC1 ({ve[0]*100:.1f}%)", fontsize=fs)
    ax.set_ylabel(f"PC2 ({ve[1]*100:.1f}%)", fontsize=fs)
    ax.set_title(title, fontsize=fs+1, fontweight="bold")
    ax.grid(True, alpha=0.15, lw=0.6)
    ax.legend(bbox_to_anchor=(1.02,1), loc="upper left", fontsize=max(fs-3,7), framealpha=0.85)
    fig.tight_layout(); return fig

def loadings_mpl(ld, top_n, title, fs=10, figsize=(9,4)):
    top = ld.sort_values("max_abs", ascending=False).head(top_n)
    x = np.arange(len(top)); w = 0.38
    fig, ax = plt.subplots(figsize=figsize)
    ax.bar(x-w/2, top["PC1"], w, label="PC1", color=TAB10[0], alpha=0.85, edgecolor="white")
    ax.bar(x+w/2, top["PC2"], w, label="PC2", color=TAB10[1], alpha=0.85, edgecolor="white")
    ax.set_xticks(x); ax.set_xticklabels(top["Feature"], rotation=45, ha="right", fontsize=fs-1)
    ax.set_ylabel("Loading", fontsize=fs); ax.set_title(title, fontsize=fs+1, fontweight="bold")
    ax.axhline(0, color="black", lw=0.7, ls="--", alpha=0.5)
    ax.legend(fontsize=fs); ax.grid(True, axis="y", alpha=0.2, lw=0.6)
    fig.tight_layout(); return fig

def unsup_scatter_mpl(Z, df_ref, pal, cl, title, xl="Dim1", yl="Dim2",
                      cluster_labels=None, fs=11, figsize=(8,6)):
    fig, ax = plt.subplots(figsize=figsize); done = set()
    for i, (_, row) in enumerate(df_ref.iterrows()):
        h = str(row[HERKUNFT_COL]); col, mkr = pal.get(h, ("#aaa","o"))
        lbl = h if h not in done else "_nolegend_"; done.add(h)
        ax.scatter(Z[i,0], Z[i,1], color=col, marker=mkr, s=60,
                   alpha=0.88, label=lbl, lw=0.4, edgecolors="white", zorder=3)
        if cluster_labels is not None:
            ax.text(Z[i,0]+0.02, Z[i,1]+0.02, str(cluster_labels[i]),
                    fontsize=7, color="#333", zorder=4)
    ax.set_xlabel(xl, fontsize=fs); ax.set_ylabel(yl, fontsize=fs)
    ax.set_title(title, fontsize=fs+1, fontweight="bold")
    ax.grid(True, alpha=0.18, lw=0.6)
    ax.legend(bbox_to_anchor=(1.02,1), loc="upper left", fontsize=max(fs-3,7), framealpha=0.85)
    fig.tight_layout(); return fig

def dendro_mpl(Z_link, labels, pal, title, fs=9, figsize=(14,5)):
    fig, ax = plt.subplots(figsize=figsize)
    dn = dendrogram(Z_link, labels=labels, ax=ax,
                    color_threshold=0, link_color_func=lambda k: "#999999",
                    leaf_font_size=fs)
    for tick, lbl in zip(ax.get_xticklabels(), dn["ivl"]):
        col, _ = pal.get(lbl, ("#555","o")); tick.set_color(col)
    ax.set_title(title, fontsize=fs+3, fontweight="bold")
    ax.set_ylabel("Distance", fontsize=fs+1); ax.grid(True, axis="y", alpha=0.2)
    fig.tight_layout(); return fig

# ─── PREPROCESSING ────────────────────────────────────────────────────────────
def coerce_num(s):
    if s.dtype == object:
        s2 = s.astype(str).str.strip().str.replace(",",".",regex=False)
        s2 = s2.replace({"":np.nan,"nan":np.nan,"None":np.nan})
        # <LOD → half LOD (extract numeric value)
        lod_mask = s2.str.startswith("<")
        lod_vals = pd.to_numeric(s2.str.lstrip("<"), errors="coerce") / 2
        s3 = pd.to_numeric(s2, errors="coerce")
        s3[lod_mask] = lod_vals[lod_mask]
        return s3
    return pd.to_numeric(s, errors="coerce")

_EC = lambda df: [c for c in [HERKUNFT_COL]+CLASS_LEVELS if c in df.columns]

def preprocess_rfa(df_in, feature_cols, zero_lod=True, lod_rule="lod2",
                   impute="median", do_log10=True, eps=1e-12,
                   scale_method="zscore"):
    df = df_in.copy(); X = df[feature_cols].copy().astype(float)
    lod_map = {}
    for c in feature_cols:
        pos = X[c][X[c] > 0]; lod_map[c] = float(pos.min()) if len(pos) else np.nan
    if zero_lod:
        for c in feature_cols:
            lod = lod_map[c]
            if np.isfinite(lod) and lod > 0:
                X.loc[X[c]==0, c] = lod/2 if lod_rule=="lod2" else lod/np.sqrt(2)
    Ximp = pd.DataFrame(SimpleImputer(strategy=impute).fit_transform(X),
                        columns=feature_cols, index=df.index)
    if do_log10:
        Ximp = np.log10(Ximp.clip(lower=float(eps)) + float(eps))
    if scale_method == "zscore":
        Xout = pd.DataFrame(StandardScaler().fit_transform(Ximp),
                             columns=feature_cols, index=df.index)
    elif scale_method == "pareto":
        std = Ximp.std(ddof=1).replace(0, 1)
        Xout = (Ximp - Ximp.mean()) / np.sqrt(std)
    else:
        Xout = Ximp
    return pd.concat([df[_EC(df)], Xout], axis=1), lod_map

def safe_ratio(d, num, den, min_denom=0.01, eps=1e-12):
    if num not in d.columns or den not in d.columns: return None
    r = d[num].astype(float) / (d[den].astype(float) + eps)
    r[d[den].astype(float).abs() < min_denom] = np.nan
    return r

def to_excel(df):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w: df.to_excel(w, index=False)
    return buf.getvalue()

def multi_excel(sheets):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        for name, df in sheets.items(): df.to_excel(w, sheet_name=name[:31], index=False)
    return buf.getvalue()

def run_pca(X, features, n=2):
    pca = PCA(n_components=n); Z = pca.fit_transform(X)
    ve = pca.explained_variance_ratio_
    ld = pd.DataFrame(pca.components_.T, index=features,
                      columns=[f"PC{i+1}" for i in range(n)])
    return Z, ve, ld, pca

# ─── STATISTICAL HELPERS ──────────────────────────────────────────────────────
def kruskal_wallis_all(df, feature_cols, class_col):
    rows = []
    for feat in feature_cols:
        groups = [g[feat].dropna().values
                  for _, g in df.groupby(class_col) if len(g[feat].dropna()) >= 2]
        if len(groups) < 2:
            rows.append({"Feature": feat, "H": np.nan, "p_value": np.nan, "Sig": "n/a"})
            continue
        try:
            H, p = sp_stats.kruskal(*groups)
            sig = "***" if p < 0.001 else ("**" if p < 0.01 else ("*" if p < 0.05 else "ns"))
            rows.append({"Feature": feat, "H": round(H,3), "p_value": round(p,4), "Sig": sig})
        except Exception:
            rows.append({"Feature": feat, "H": np.nan, "p_value": np.nan, "Sig": "err"})
    return pd.DataFrame(rows).sort_values("p_value")

def hotelling_t2(Z_train, Z_any, alpha=0.05):
    from scipy.stats import f as f_dist
    n, p = Z_train.shape
    mu = Z_train.mean(axis=0)
    S = np.cov(Z_train.T) if p > 1 else np.array([[np.var(Z_train[:,0])]])
    try:
        S_inv = np.linalg.pinv(S)
    except Exception:
        S_inv = np.eye(p)
    T2 = np.array([(z-mu) @ S_inv @ (z-mu) for z in Z_any])
    if n > p:
        T2_crit = p*(n-1)/(n-p) * f_dist.ppf(1-alpha, p, n-p)
    else:
        T2_crit = np.percentile(T2, 95)
    return T2, float(T2_crit)

# ─── MODELZOO HELPERS ─────────────────────────────────────────────────────────
def mz_lod_log(df_raw, features, lod_map, lod_rule="lod2", do_log10=True, eps=1e-12):
    X = df_raw[features].copy().astype(float)
    if lod_map:
        for c in features:
            lod = lod_map.get(c, np.nan)
            if np.isfinite(lod) and lod > 0:
                X.loc[X[c]==0, c] = lod/2 if lod_rule=="lod2" else lod/np.sqrt(2)
    if do_log10:
        X = np.log10(X.clip(lower=float(eps)) + float(eps))
    return X

def _prep_pipe(scale_method="zscore"):
    sc = StandardScaler() if scale_method == "zscore" else "passthrough"
    return Pipeline([("imp", SimpleImputer(strategy="median")), ("sc", sc)])

def mz_build_pipes(seed, knn_k, svm_c, rf_trees, scale_method="zscore"):
    prep = _prep_pipe(scale_method)
    models = {
        "LDA": LDA(),
        f"kNN(k={knn_k})": KNeighborsClassifier(n_neighbors=int(knn_k)),
        f"SVM(C={svm_c:g})": SVC(kernel="rbf", C=float(svm_c), gamma="scale",
                                  probability=True, random_state=int(seed)),
        f"RF({rf_trees})": RandomForestClassifier(n_estimators=int(rf_trees),
                             class_weight="balanced", random_state=int(seed), n_jobs=-1),
    }
    return {n: Pipeline([("prep", clone(prep)), ("clf", m)]) for n, m in models.items()}

def mz_kfold(X, y, pipes, k=5, seed=42):
    skf = StratifiedKFold(n_splits=k, shuffle=True, random_state=seed)
    rows = []
    for name, pipe in pipes.items():
        cv = cross_validate(clone(pipe), X, y, cv=skf,
                            scoring={"acc":"accuracy",
                                     "balacc":"balanced_accuracy",
                                     "f1":"f1_macro"},
                            return_train_score=False, n_jobs=-1)
        rows.append({"Model": name,
                     f"CV{k}_Acc_mean":    round(cv["test_acc"].mean(),4),
                     f"CV{k}_Acc_std":     round(cv["test_acc"].std(),4),
                     f"CV{k}_BalAcc_mean": round(cv["test_balacc"].mean(),4),
                     f"CV{k}_F1_mean":     round(cv["test_f1"].mean(),4),
                     f"CV{k}_F1_std":      round(cv["test_f1"].std(),4)})
    return pd.DataFrame(rows).sort_values(f"CV{k}_F1_mean", ascending=False).reset_index(drop=True)

def mz_evaluate(df_num, label_col, features, test_size, seed,
                knn_k, svm_c, rf_trees, lod_map=None, lod_rule="lod2",
                do_log10=True, eps=1e-12, kfold_k=5, scale_method="zscore"):
    X = mz_lod_log(df_num, features, lod_map, lod_rule, do_log10, eps)
    y = df_num[label_col].astype(str)
    Xtr, Xte, ytr, yte = train_test_split(X, y, test_size=float(test_size),
                                           stratify=y, random_state=int(seed))
    pipes = mz_build_pipes(seed, knn_k, svm_c, rf_trees, scale_method)
    # k-fold CV (primary)
    cv_res = mz_kfold(X, y, pipes, k=kfold_k, seed=seed)
    # Train/test split (secondary)
    rows, confmats, preds = [], {}, {}
    lbls = sorted(y.unique())
    for name, pipe in pipes.items():
        pipe.fit(Xtr, ytr); pred = pipe.predict(Xte)
        rows.append({"Model": name,
                     "Split_Acc":    round(accuracy_score(yte, pred),4),
                     "Split_BalAcc": round(balanced_accuracy_score(yte, pred),4),
                     "Split_F1":     round(f1_score(yte, pred, average="macro"),4),
                     "n_train": len(ytr), "n_test": len(yte)})
        confmats[name] = (lbls, confusion_matrix(yte, pred, labels=lbls))
        preds[name] = pred
    split_res = (pd.DataFrame(rows).sort_values("Split_F1", ascending=False)
                 .reset_index(drop=True))
    return dict(features=features, X_train=Xtr, X_test=Xte,
                y_train=ytr, y_test=yte,
                pipes=pipes, cv_results=cv_res, split_results=split_res,
                confmats=confmats, preds_test=preds,
                df_train=df_num, label_col=label_col,
                lod_map=lod_map, lod_rule=lod_rule,
                do_log10=do_log10, eps=eps,
                seed=seed, test_size=test_size,
                knn_k=knn_k, svm_c=svm_c, rf_trees=rf_trees,
                kfold_k=kfold_k, scale_method=scale_method)

def mz_q_res(prep, X_train, X_any, n=2):
    Xtr = prep.fit_transform(X_train); pca = PCA(n_components=int(n))
    Ztr = pca.fit_transform(Xtr)
    Qtr = np.sum((Xtr - pca.inverse_transform(Ztr))**2, axis=1)
    Xa = prep.transform(X_any); Za = pca.transform(Xa)
    return Qtr, np.sum((Xa - pca.inverse_transform(Za))**2, axis=1), Ztr, Za, pca

def mz_pca_proj(pipe, Xtr, Xte, Xunk=None):
    prep = pipe.named_steps["prep"]
    Ptr = prep.fit_transform(Xtr); Pte = prep.transform(Xte)
    pca = PCA(n_components=2); Ztr = pca.fit_transform(Ptr); Zte = pca.transform(Pte)
    Zunk = pca.transform(prep.transform(Xunk)) if Xunk is not None else None
    return Ztr, Zte, Zunk, pca.explained_variance_ratio_

def mz_hybrid(state, model_name, df_unk, name_col, prob_thr, q_quant, q_pca, topk):
    pipe = state["pipes"][model_name]; features = state["features"]
    Xunk = mz_lod_log(df_unk, features, state.get("lod_map"),
                       state.get("lod_rule","lod2"), state.get("do_log10",True),
                       state.get("eps",1e-12))
    Qtr, Qunk, Ztr, Zunk_sc, _ = mz_q_res(pipe.named_steps["prep"],
                                             state["X_train"], Xunk, n=q_pca)
    _, Ztr_pca, _, _, _ = mz_q_res(pipe.named_steps["prep"],
                                    state["X_train"], state["X_train"], n=q_pca)
    T2unk, T2crit = hotelling_t2(Ztr_pca, Zunk_sc)
    qthr = float(np.quantile(Qtr, float(q_quant)))
    clf = pipe.named_steps["clf"]; proba = pipe.predict_proba(Xunk)
    classes = clf.classes_; tidx = np.argsort(-proba, axis=1)[:, :int(topk)]
    out = pd.DataFrame(index=df_unk.index)
    out.insert(0, "Bauteil",
               df_unk[name_col].values if (name_col and name_col in df_unk.columns)
               else [f"Bauteil_{i+1}" for i in range(len(df_unk))])
    for k in range(int(topk)):
        out[f"Top{k+1}_Class"] = classes[tidx[:,k]]
        out[f"Top{k+1}_Prob"]  = proba[np.arange(len(proba)), tidx[:,k]]
    out["Q_residual"] = Qunk; out["Q_threshold"] = qthr
    out["T2"] = T2unk; out["T2_threshold"] = T2crit
    out["Pass_Q"]    = Qunk <= qthr
    out["Pass_T2"]   = T2unk <= T2crit
    out["Pass_Prob"] = out["Top1_Prob"] >= float(prob_thr)
    out["Assigned"]  = out["Pass_Q"] & out["Pass_T2"] & out["Pass_Prob"]
    # Explicit reason for UNASSIGNED
    def reason(row):
        if row["Assigned"]: return "✅ Assigned"
        r = []
        if not row["Pass_Q"]:    r.append("Q-Residual too high")
        if not row["Pass_T2"]:   r.append("Hotelling T² too high")
        if not row["Pass_Prob"]: r.append("P(Class) too low")
        return " + ".join(r)
    out["Grund"] = out.apply(reason, axis=1)
    out["Predicted"] = np.where(out["Assigned"], out["Top1_Class"], "UNASSIGNED")
    s = {"n":len(out), "assigned":int(out["Assigned"].sum()),
         "unassigned":int((~out["Assigned"]).sum()),
         "pass_Q":int(out["Pass_Q"].sum()), "pass_T2":int(out["Pass_T2"].sum()),
         "pass_Prob":int(out["Pass_Prob"].sum()),
         "q_threshold":qthr, "T2_threshold":T2crit}
    return out.reset_index(drop=True), s

def mz_thr_curve(state, model_name, thresholds, q_quant, q_pca):
    pipe = state["pipes"][model_name]
    Xtr = state["X_train"]; Xte = state["X_test"]
    yte = state["y_test"].astype(str).values
    proba = pipe.predict_proba(Xte); classes = pipe.named_steps["clf"].classes_
    top1c = classes[np.argmax(proba, axis=1)].astype(str)
    top1p = proba[np.arange(len(proba)), np.argmax(proba, axis=1)]
    Qtr, Qte, Ztr, Zte, _ = mz_q_res(pipe.named_steps["prep"], Xtr, Xte, n=q_pca)
    T2te, T2crit = hotelling_t2(Ztr, Zte)
    qthr = float(np.quantile(Qtr, float(q_quant)))
    rows = []
    for thr in thresholds:
        thr = float(thr)
        asgn = (Qte <= qthr) & (T2te <= T2crit) & (top1p >= thr)
        cov = float(asgn.mean()); acc = mf1 = np.nan
        if asgn.sum():
            acc = accuracy_score(yte[asgn], top1c[asgn])
            mf1 = f1_score(yte[asgn], top1c[asgn], average="macro", zero_division=0)
        rows.append({"prob_threshold":thr, "coverage":cov,
                     "acc_assigned":acc, "macroF1_assigned":mf1})
    return pd.DataFrame(rows)

def mz_repeated(state, repeats, seed_base):
    df_num = state["df_train"]; label_col = state["label_col"]
    features = state["features"]; test_size = state["test_size"]
    lod_map = state.get("lod_map"); lod_rule = state.get("lod_rule","lod2")
    do_log10 = state.get("do_log10",True); eps = state.get("eps",1e-12)
    knn_k = state["knn_k"]; svm_c = state["svm_c"]; rf_trees = state["rf_trees"]
    scale_method = state.get("scale_method","zscore")
    X = mz_lod_log(df_num, features, lod_map, lod_rule, do_log10, eps)
    y = df_num[label_col].astype(str).values
    rows = []
    for r in range(int(repeats)):
        rs = int(seed_base) + r
        pipes = mz_build_pipes(rs, knn_k, svm_c, rf_trees, scale_method)
        Xtr, Xte, ytr, yte = train_test_split(X, y, test_size=float(test_size),
                                               random_state=rs, stratify=y)
        for name, pipe in pipes.items():
            p = clone(pipe); p.fit(Xtr, ytr); pred = p.predict(Xte)
            rows.append({"repeat":r, "model":name,
                         "accuracy":accuracy_score(yte,pred),
                         "balanced_accuracy":balanced_accuracy_score(yte,pred),
                         "macroF1":f1_score(yte,pred,average="macro",zero_division=0)})
    df_r = pd.DataFrame(rows)
    agg = (df_r.groupby("model")[["accuracy","balanced_accuracy","macroF1"]]
           .agg(["mean","std"]).reset_index())
    agg.columns = ["model"] + [f"{m}_{s}" for m,s in agg.columns.tolist()[1:]]
    return df_r, agg.sort_values("macroF1_mean",ascending=False).reset_index(drop=True)

def mz_permutation(state, model_name, n_perm=99, cv_k=5):
    pipe  = state["pipes"][model_name]; features = state["features"]
    df_num = state["df_train"]; label_col = state["label_col"]
    X = mz_lod_log(df_num, features, state.get("lod_map"),
                   state.get("lod_rule","lod2"), state.get("do_log10",True),
                   state.get("eps",1e-12))
    y = df_num[label_col].astype(str).values
    skf = StratifiedKFold(n_splits=cv_k, shuffle=True, random_state=state["seed"])
    score, perm_scores, pvalue = permutation_test_score(
        clone(pipe), X, y, scoring="balanced_accuracy",
        cv=skf, n_permutations=n_perm,
        random_state=state["seed"], n_jobs=-1)
    return score, perm_scores, pvalue

def mz_feature_importance(state, model_name):
    pipe = state["pipes"][model_name]; features = state["features"]
    clf = pipe.named_steps["clf"]
    fi_df = None
    if hasattr(clf, "feature_importances_"):
        fi = clf.feature_importances_
        fi_df = pd.DataFrame({"Feature":features,"Importance":fi})
        fi_df = fi_df.sort_values("Importance",ascending=False).reset_index(drop=True)
    # Also compute coef for LDA
    elif hasattr(clf, "coef_"):
        fi = np.abs(clf.coef_).mean(axis=0)
        fi_df = pd.DataFrame({"Feature":features,"Importance":fi/fi.sum()})
        fi_df = fi_df.sort_values("Importance",ascending=False).reset_index(drop=True)
    return fi_df

# ─── METHODS DOCX GENERATOR ───────────────────────────────────────────────────
def generate_methods_docx():
    """Generate a Word document with chemometric method explanations. Returns bytes."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"; style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    for lv, sz in [(1,16),(2,13)]:
        h = doc.styles[f"Heading {lv}"]
        h.font.name = "Calibri"; h.font.size = Pt(sz); h.font.bold = True
        h.font.color.rgb = RGBColor(0x1A,0x1F,0x2E)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run("Methodological Background\nChemometric Analysis for Provenance Determination")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1A,0x1F,0x2E)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp.space_after = Pt(20)
    sr = sp.add_run("Supplementary methods documentation for elemental provenance analysis")
    sr.italic = True; sr.font.size = Pt(10); sr.font.color.rgb = RGBColor(0x66,0x66,0x66)

    _SECS = [
      ("1. Principal Component Analysis (PCA)",
       "PCA is an unsupervised dimensionality reduction technique that projects high-dimensional "
       "data onto orthogonal components ordered by decreasing variance. Scores show sample "
       "positions; loadings indicate the contribution of original variables to each PC.",
       "Principal Component Analysis (PCA) is one of the most widely used unsupervised multivariate "
       "methods in chemometrics and geochemical provenance studies. It transforms a set of possibly "
       "correlated variables into a smaller number of uncorrelated principal components (PCs) by "
       "performing an eigendecomposition of the covariance matrix. The first PC captures the "
       "direction of maximum variance in the data, and each subsequent PC is orthogonal to the "
       "previous ones while explaining decreasing amounts of remaining variance.\n\n"
       "In the context of mineral analysis, PCA serves as a critical exploratory tool. The scores "
       "plot projects samples into the reduced PC space, enabling visual assessment of natural "
       "groupings or outliers related to provenance, mineral class, or geological setting. The "
       "loadings indicate which elemental variables (e.g., Rb, K, Cs, Fe, Mg) contribute most to "
       "each component, thereby revealing the geochemical signatures that drive sample separation. "
       "A biplot combines both scores and loadings in a single visualization.\n\n"
       "PCA requires appropriate preprocessing, typically log-transformation and autoscaling "
       "(z-score normalization), to ensure that variables measured on different scales contribute "
       "equally. The number of meaningful components is typically determined by the Kaiser criterion "
       "(eigenvalue > 1) or by inspecting the cumulative explained variance."),
      ("2. Partial Least Squares Discriminant Analysis (PLS-DA)",
       "PLS-DA is a supervised classification method that identifies latent variables maximizing "
       "the covariance between predictor matrix X and a dummy-coded class matrix Y. The VIP score "
       "quantifies each variable's contribution to the discrimination.",
       "Partial Least Squares Discriminant Analysis (PLS-DA) combines the dimensionality reduction "
       "properties of PLS regression with class discrimination. Unlike PCA, which maximizes variance "
       "in X alone, PLS-DA identifies latent variables (LVs) that maximize the covariance between "
       "the predictor matrix X and a dummy-coded response matrix Y representing class membership.\n\n"
       "The method is particularly suitable for chemometric datasets where the number of variables "
       "approaches or exceeds the number of observations, and where predictors are strongly "
       "intercorrelated — both common characteristics of multi-elemental analytical data. A key diagnostic tool "
       "is the Variable Importance in Projection (VIP) score, which quantifies the contribution of "
       "each original variable to the model. Variables with VIP > 1 are generally considered "
       "important discriminators.\n\n"
       "Model validation is performed through stratified k-fold cross-validation and permutation "
       "testing to assess whether the observed classification accuracy exceeds chance level. When "
       "the number of classes is small (e.g., two mineral classes), PLS-DA may extract only a single "
       "latent variable. In such cases, the scores plot reduces to a one-dimensional representation."),
      ("3. Clustering Methods",
       "k-Means partitions samples into k clusters by minimizing within-cluster sum of squares. "
       "The Silhouette score measures cluster cohesion vs. separation. Hierarchical clustering "
       "builds a dendrogram reflecting nested groupings.",
       "Clustering methods are unsupervised techniques that group samples based on similarity "
       "without prior class labels. k-Means clustering assigns n observations to k clusters by "
       "iteratively minimizing the within-cluster sum of squares (WCSS). The Elbow method plots "
       "WCSS against cluster count to identify the point of diminishing returns. The Silhouette "
       "score (ranging from -1 to +1) provides a quantitative measure of cluster quality: values "
       "above 0.5 indicate well-separated clusters.\n\n"
       "Hierarchical agglomerative clustering builds a tree-like dendrogram by iteratively merging "
       "the most similar pairs of clusters. Various linkage criteria (Ward, complete, average) and "
       "distance metrics (Euclidean, cosine) control the merging behavior. Ward's method minimizes "
       "the increase in total within-cluster variance at each step.\n\n"
       "In mineral provenance analysis, clustering serves as an independent validation of groups "
       "identified by PCA. Consistent results across multiple methods strengthen the evidence for "
       "geochemically distinct provenance groups."),
      ("4. Descriptive Statistics",
       "Descriptive statistics summarize central tendency, dispersion, and distribution shape. "
       "The Kruskal-Wallis test assesses whether group distributions differ significantly (p < 0.05).",
       "Descriptive statistics provide a foundational summary of the dataset prior to multivariate "
       "analysis. Key measures include arithmetic mean, median, standard deviation, minimum, maximum, "
       "and interquartile range. Graphical representations — boxplots, violin plots, histograms with "
       "kernel density estimation (KDE), and correlation matrices — facilitate visual assessment of "
       "distributions, outliers, and inter-element relationships.\n\n"
       "For group comparisons, the Kruskal-Wallis H-test is employed as a non-parametric alternative "
       "to one-way ANOVA. It tests whether independent samples from different provenance classes "
       "originate from the same distribution. Significance levels (p < 0.05, 0.01, 0.001) indicate "
       "increasingly strong evidence that at least one group differs. Elements with highly significant "
       "results are candidates for discriminating markers in subsequent supervised analyses."),
      ("5. Data Preprocessing",
       "Preprocessing includes LOD treatment, imputation, log-transformation, and scaling (z-score "
       "or Pareto). These steps reduce skewness, equalize variable scales, and stabilize variance.",
       "Data preprocessing ensures the quality and comparability of measured variables before "
       "multivariate analysis. The pipeline typically consists of: (1) Below-LOD treatment — values "
       "below the limit of detection are replaced by LOD/2 or LOD/sqrt(2). (2) Imputation — remaining "
       "missing values are replaced using the column median. (3) Logarithmic transformation (log10) "
       "— reduces skewness, stabilizes variance, and converts multiplicative relationships into "
       "additive ones. (4) Scaling — z-score normalization (autoscaling) subtracts the mean and "
       "divides by the standard deviation, giving each variable unit variance. Pareto scaling divides "
       "by the square root of the standard deviation, providing a less aggressive alternative.\n\n"
       "The choice of scaling method affects which variables dominate the multivariate model and "
       "should be reported explicitly in any scientific publication."),
      ("6. Feature Selection and Variable Importance",
       "Feature importance is assessed via PCA loadings, PLS-DA VIP scores, and Random Forest "
       "feature importances. Kruskal-Wallis tests provide univariate pre-screening.",
       "Feature selection serves to identify the subset of measured elements that contribute most "
       "to provenance discrimination. PCA loadings reveal which elements drive variance along each "
       "principal component. PLS-DA VIP scores quantify the weighted contribution of each variable "
       "to class discrimination — VIP > 1 is considered significant, below 0.5 is negligible. "
       "Random Forest models provide a complementary assessment through Gini importance or "
       "permutation importance, both robust to nonlinear relationships.\n\n"
       "Combining multiple feature importance metrics from different methods increases confidence "
       "in the identified marker elements. For mineral provenance analysis, elements such as Rb, K, "
       "Cs, Fe, Mg, and Ti are frequently identified as key discriminators due to their sensitivity "
       "to geological formation conditions."),
    ]

    for title, short, extended in _SECS:
        doc.add_heading(title, level=1)
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(short)
        doc.add_heading("Extended Description", level=2)
        for para in extended.split("\n\n"):
            doc.add_paragraph(para.strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# ─── PLS-DA HELPERS ───────────────────────────────────────────────────────────
def compute_vip(pls_model):
    T = pls_model.x_scores_
    W = pls_model.x_weights_
    Q = pls_model.y_loadings_
    p = W.shape[0]; n_comp = T.shape[1]
    SS = np.array([float(T[:,h].T @ T[:,h]) * float(Q[:,h].T @ Q[:,h])
                   for h in range(n_comp)])
    total_SS = SS.sum()
    if total_SS == 0:
        return np.zeros(p)
    W_norm = W / (np.linalg.norm(W, axis=0) + 1e-12)
    VIP = np.sqrt(p * (W_norm**2 @ SS) / total_SS)
    return VIP

def fit_plsda(X_train, y_train, n_components=2):
    classes = sorted(np.unique(y_train))
    n_comp = min(n_components, len(classes)-1 if len(classes) > 1 else 1,
                 X_train.shape[1], X_train.shape[0]-1)
    n_comp = max(n_comp, 1)
    Y = label_binarize(y_train, classes=classes)
    if Y.shape[1] == 1:
        Y = np.hstack([1-Y, Y])
    pls = PLSRegression(n_components=n_comp, max_iter=1000)
    pls.fit(X_train, Y)
    return pls, classes

def predict_plsda(pls, classes, X):
    Y_pred = pls.predict(X)
    idx = np.argmax(Y_pred, axis=1)
    idx = np.clip(idx, 0, len(classes)-1)
    return np.array(classes)[idx]

def plsda_kfold(X, y, n_components=2, k=5, seed=42):
    skf = StratifiedKFold(n_splits=k, shuffle=True, random_state=seed)
    accs, f1s = [], []
    for tr_idx, te_idx in skf.split(X, y):
        pls, cls = fit_plsda(X[tr_idx], y[tr_idx], n_components)
        ypred = predict_plsda(pls, cls, X[te_idx])
        accs.append(accuracy_score(y[te_idx], ypred))
        f1s.append(f1_score(y[te_idx], ypred, average="macro", zero_division=0))
    return np.array(accs), np.array(f1s)

def check_class_sizes(df, label_col, min_samples=3):
    counts = df[label_col].astype(str).value_counts()
    small = counts[counts < min_samples]
    return counts, small

# ─── NAVIGATION ───────────────────────────────────────────────────────────────
SEC_HOME="🏠 Home"; SEC_EXP="🔬 Exploration"; SEC_MZ="🤖 Model Zoo"
EXP_PAGES = [
    "📂 1.1 · Data & Classes",
    "✂️ 1.2 · Features",
    "⚙️ 1.3 · Preprocessing",
    "📊 1.4 · Descriptive Statistics",
    "🎨 1.5 · Color Scheme",
    "🔬 1.6 · Screening",
    "➗ 1.7 · Ratios",
    "📈 1.8 · PCA Global",
    "🧩 1.9 · Sample Groups",
    "📉 1.10 · Univariate Scatter",
]
SUP_PAGES = [
    "⚙️ S.1 · Training & k-Fold CV",
    "🔬 S.2 · PLS-DA",
    "📊 S.3 · Results",
    "🗺️ S.4 · PCA Space",
    "🚪 S.5 · Outlier & Gate",
    "📉 S.6 · Threshold",
    "🔁 S.7 · Validation & Permutation",
    "📤 S.8 · Export",
]
UNS_PAGES = [
    "🌀 U.1 · PCA Biplot",
    "🔵 U.2 · k-Means",
    "🌲 U.3 · Dendrogram",
    "🗺️ U.4 · UMAP",
]

st.sidebar.markdown("""
<div style="text-align: center; padding: 0.5rem 0 0.8rem 0;">
    <span style="font-size: 1.6rem;">🔬</span><br>
    <span style="font-size: 1.1rem; font-weight: 700; letter-spacing: 0.04em;
                 color: #e8ecf1;">ChemoTrace</span><br>
    <span style="font-size: 0.72rem; color: #8899aa; letter-spacing: 0.06em;
                 text-transform: uppercase;">Multivariate Provenance Analysis</span>
</div>
""", unsafe_allow_html=True)
_section = st.sidebar.radio("Bereich", [SEC_HOME,SEC_EXP,SEC_MZ], label_visibility="collapsed")
_page = None; _mz_mode = None
if _section == SEC_EXP:
    _page = st.sidebar.radio("Schritt", EXP_PAGES, label_visibility="collapsed")
elif _section == SEC_MZ:
    _mz_mode = st.sidebar.radio("Modus", ["🧑‍🏫 Supervised","🔍 Unsupervised"], label_visibility="collapsed")
    _page = st.sidebar.radio("Schritt",
                               SUP_PAGES if "Supervised" in _mz_mode else UNS_PAGES,
                               label_visibility="collapsed")

st.sidebar.markdown("---")
st.sidebar.markdown("### 📁 Data Source")
_ds_options = ["📤 Upload own file", "📦 Use demo dataset"]
_ds_idx = 0 if st.session_state.data_source == "upload" else 1
_ds_choice = st.sidebar.radio("Data Source", _ds_options, index=_ds_idx,
                               label_visibility="collapsed")
_new_source = "upload" if "Upload" in _ds_choice else "demo"
if _new_source != st.session_state.data_source:
    # Source changed → reset all loaded data
    st.session_state.data_source = _new_source
    st.session_state.demo_loaded = False
    for _rk in ["df_raw","df_klassen","df_merged","df_num","df_proc",
                 "feature_cols","lod_map","df_ratios","df_grouped",
                 "pca_scores","pca_loadings","pca_group_scores",
                 "pca_group_loadings","mz_state","mz_df_unknown","mz_unk_name_col"]:
        st.session_state[_rk] = None
    st.rerun()
if st.session_state.data_source == "demo":
    st.sidebar.caption("🟡 Demo mode active")
else:
    st.sidebar.caption("Upload Excel/CSV file in step 1.1")

st.sidebar.markdown("---")
st.sidebar.markdown("### Status")
_dot = lambda c: "🟢" if c else "🔴"
st.sidebar.markdown(f"{_dot(st.session_state.df_raw is not None)} Data loaded")
st.sidebar.markdown(f"{_dot(st.session_state.df_merged is not None)} Classes joined")
st.sidebar.markdown(f"{_dot(st.session_state.df_proc is not None)} Preprocessing ✓")
st.sidebar.markdown(f"{_dot(st.session_state.mz_state is not None)} Models trained")

st.sidebar.markdown("---")
if st.session_state.df_merged is not None:
    st.sidebar.markdown("### 🎨 Visualization")
    _cl = st.sidebar.radio("Class Level", CLASS_LEVELS,
                            index=CLASS_LEVELS.index(st.session_state.class_level),
                            label_visibility="collapsed")
    st.session_state.class_level = _cl
    st.session_state.symbol_detail = st.sidebar.checkbox(
        "🔠 Symbol = Sample Origin", value=st.session_state.symbol_detail)

st.sidebar.markdown("---")
if st.sidebar.button("🔒 Logout", use_container_width=True):
    st.session_state.authenticated = False
    st.rerun()

# ══════════════════════════════════════════════════════════════════════════════
# STARTSEITE
# ══════════════════════════════════════════════════════════════════════════════
if _section == SEC_HOME:
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #1a1f2e 0%, #2c3e6b 60%, #4a6fa5 100%);
        border-radius: 12px;
        padding: 2.5rem 2rem 1.8rem 2rem;
        margin-bottom: 1.5rem;
        box-shadow: 0 4px 20px rgba(26, 31, 46, 0.15);
    ">
        <h1 style="color: #e8ecf1; margin: 0 0 0.3rem 0; font-size: 2.2rem;
                    font-weight: 700; letter-spacing: -0.01em; border: none; padding: 0;">
            🔬 ChemoTrace
        </h1>
        <p style="color: #b8c4d0; margin: 0 0 0.6rem 0; font-size: 1.05rem; letter-spacing: 0.02em;">
            Multivariate Chemometric Analysis for Provenance Determination
        </p>
        <p style="color: #8899aa; margin: 0; font-size: 0.78rem; letter-spacing: 0.03em;">
            University of Augsburg · Chair of Resource and Chemical Engineering ·
            Institute of Materials Resource Management · In cooperation with BMW Group
        </p>
    </div>
    """, unsafe_allow_html=True)

    # ── Scientific scope description ──────────────────────────────
    st.markdown("""
    <div style="
        background: linear-gradient(135deg, #f4f6f9 0%, #e8ecf1 100%);
        border-left: 4px solid #4a6fa5;
        border-radius: 6px;
        padding: 1rem 1.2rem;
        margin-bottom: 1.2rem;
        font-size: 0.92rem;
        color: #2c3e6b;
    ">
        <strong>Scientific Scope</strong> — Multivariate statistical methods are used to identify patterns
        in compositional data and assign samples to potential origins based on similarity structures.
        This tool supports exploratory analysis, supervised classification, and unsupervised clustering
        for provenance determination of minerals and materials based on elemental fingerprints.
    </div>
    """, unsafe_allow_html=True)

    c1, c2 = st.columns(2, gap="large")
    with c1:
        st.markdown("""
## 🔬 Exploration
1. 📂 Data Import (measurements + classes — Fuzzy-Join)
2. ✂️ Feature Selection
3. ⚙️ Preprocessing (LOD, log₁₀, z/Pareto scaling)
4. **📊 Descriptive Statistics**
   - Summary tables, boxplots, histograms
   - Correlation matrix, concentration heatmap
   - Kruskal-Wallis significance tests, pairplot
5. 🎨 Color Scheme (class hierarchy)
6. 🔬 Screening Scatter
7. ➗ Element Ratios
8. 📈 PCA Global + Biplot
9. 🧩 Sample Groups (on raw data)
10. 📉 Univariate Scatter
        """)
    with c2:
        st.markdown("""
## 🤖 Model Zoo
### 🧑‍🏫 Supervised
- S.1 Training with **k-Fold CV** as primary metric
- **S.2 PLS-DA** + VIP Scores
- S.3 Results: norm. confusion matrix, Feature Importance
- S.4 PCA Space (Train + Test + Unknowns)
- S.5 Outlier Gate: Q-Residual + **Hotelling T²** + rejection logic
- S.6 Threshold Analysis
- S.7 Repeated Validation + **Permutation Test**
- S.8 Complete Export

### 🔍 Unsupervised
PCA Biplot · k-Means (Silhouette) · Dendrogram · UMAP
        """)
    st.info("**Class Hierarchy:** Class Level → Country → ID → Sample Origin (Symbol)\n\n"
            "Color = active class level (adjustable in sidebar)\n\n"
            "Symbol = individual sample origin (toggle in sidebar)")

    # ── Methods documentation download ─────────────────────────────
    st.divider()
    st.markdown("### 📄 Methods Documentation")
    _docx_bytes = generate_methods_docx()
    if _docx_bytes:
        st.download_button(
            "📥 Methodological Background (Word)",
            _docx_bytes,
            "chemometrics_methods.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=False,
        )
        st.caption("Scientific descriptions of PCA, PLS-DA, Clustering, "
                   "Descriptive Statistics, Preprocessing, and Feature Selection.")
    else:
        st.warning("⚠️ `python-docx` not installed — Word export unavailable.")

# ══════════════════════════════════════════════════════════════════════════════
# SEITE 1 — EXPLORATION
# ══════════════════════════════════════════════════════════════════════════════
elif _section == SEC_EXP:

    # 1.1 ──────────────────────────────────────────────────────────────────────
    if _page == EXP_PAGES[0]:
        st.title("📂 1.1 · Data Import & Class Join")

        # ── Helper: process & join a loaded Excel/CSV into session state ──
        def _process_and_join(df_d, df_k, source_label="Datei"):
            """Shared logic for both upload and demo paths."""
            dm = df_d.copy()
            match_report = pd.DataFrame()
            if df_k is not None:
                dm, match_report = join_klassen(dm, df_k)
            else:
                for cl2 in CLASS_LEVELS:
                    dm[cl2] = "n/a"
            pf = [c for c in dm.columns if c not in [HERKUNFT_COL] + CLASS_LEVELS]
            for c in pf:
                dm[c] = coerce_num(dm[c])
            fc = [c for c in pf if np.issubdtype(dm[c].dtype, np.number)]
            for k2 in ["df_num", "df_proc", "df_ratios", "df_grouped", "lod_map",
                        "pca_scores", "pca_loadings", "pca_group_scores",
                        "pca_group_loadings", "mz_state"]:
                st.session_state[k2] = None
            st.session_state.df_raw = df_d
            st.session_state.df_klassen = df_k
            st.session_state.df_merged = dm
            st.session_state.df_num = dm
            st.session_state.feature_cols = fc
            st.success(f"✅ {dm.shape[0]} samples · {len(fc)} Features  ({source_label})")
            if len(match_report):
                st.subheader("🔍 Fuzzy-Matching Control Table")
                st.dataframe(match_report.style.map(
                    lambda v: "color:red" if v == "❌ kein Match" else
                    ("color:orange" if "fuzzy" in str(v) else "color:green"),
                    subset=["Match"]), use_container_width=True)
                unm = match_report[match_report["Match"] == "❌ kein Match"]["Daten"].tolist()
                if unm:
                    st.warning(f"⚠️ No match found: {unm}")
                else:
                    st.success("✅ All sample origins matched successfully!")

        # ── DEMO MODE ────────────────────────────────────────────────────
        if st.session_state.data_source == "demo":
            st.info("📦 **Demo Mode** — The built-in demo dataset is being used.  \n"
                    "Switch to *Upload own file* in the sidebar to use your own data.")
            if not DEMO_FILE.exists():
                st.error(f"⚠️ Demo file not found: `{DEMO_FILE}`\n\n"
                         "Please place an Excel file at `data/demo.xlsx` im Repository ab "
                         "(with sheets *Daten* and *Klassen*).")
                st.stop()

            # Auto-load demo on first visit or after source switch
            if not st.session_state.demo_loaded or st.session_state.df_merged is None:
                try:
                    xls = pd.ExcelFile(DEMO_FILE)
                    d_opt = [s for s in xls.sheet_names if "daten" in s.lower()]
                    k_opt = [s for s in xls.sheet_names if "klass" in s.lower()]
                    sh_d = d_opt[0] if d_opt else xls.sheet_names[0]
                    sh_k = k_opt[0] if k_opt else (xls.sheet_names[1] if len(xls.sheet_names) > 1 else None)
                    df_d = pd.read_excel(xls, sheet_name=sh_d)
                    df_k = pd.read_excel(xls, sheet_name=sh_k) if sh_k else None
                    _process_and_join(df_d, df_k, source_label="Demo Dataset")
                    st.session_state.demo_loaded = True
                except Exception as e:
                    st.error(f"Error loading demo file: {e}")
            else:
                dm2 = st.session_state.df_merged
                st.success(f"✅ {dm2.shape[0]} samples · {len(st.session_state.feature_cols)} Features  (Demo Dataset)")
                st.dataframe(dm2[[HERKUNFT_COL] + CLASS_LEVELS].drop_duplicates(), use_container_width=True)

        # ── UPLOAD MODE (original logic, preserved) ──────────────────────
        else:
            st.info("Upload Excel with **Data** sheet (measurements) and **Classes** sheet (hierarchy).\n"
                    "The app joins both automatically via fuzzy-matching and displays a matching table.")
            uploaded = st.file_uploader("Excel / CSV", type=["xlsx", "xls", "csv", "txt"])
            if uploaded:
                try:
                    ext = uploaded.name.split(".")[-1].lower()
                    if ext in ("xlsx", "xls"):
                        xls = pd.ExcelFile(uploaded)
                        st.markdown(f"**Sheets:** {', '.join(xls.sheet_names)}")
                        c1, c2 = st.columns(2)
                        d_opt = [s for s in xls.sheet_names if "daten" in s.lower()]
                        k_opt = [s for s in xls.sheet_names if "klass" in s.lower()]
                        with c1:
                            sh_d = st.selectbox("Data Sheet", xls.sheet_names,
                                                index=xls.sheet_names.index(d_opt[0]) if d_opt else 0)
                        with c2:
                            sh_k = st.selectbox("Class Sheet", xls.sheet_names,
                                                index=xls.sheet_names.index(k_opt[0]) if k_opt else min(1, len(xls.sheet_names) - 1))
                        df_d = pd.read_excel(xls, sheet_name=sh_d)
                        df_k = pd.read_excel(xls, sheet_name=sh_k)
                    else:
                        sep = st.selectbox("Separator", [",", ";", "\t"])
                        df_d = pd.read_csv(uploaded, sep=sep)
                        df_k = None
                        st.warning("CSV detected — no class sheet available. Please add manually or use Excel format.")

                    if st.button("🔗 Join & Apply", type="primary"):
                        _process_and_join(df_d, df_k, source_label="Upload")

                    elif st.session_state.df_merged is not None:
                        dm2 = st.session_state.df_merged
                        st.dataframe(dm2[[HERKUNFT_COL] + CLASS_LEVELS].drop_duplicates(), use_container_width=True)
                except Exception as e:
                    st.error(f"Error: {e}")
            elif st.session_state.df_merged is not None:
                dm2 = st.session_state.df_merged
                st.success(f"✅ {dm2.shape[0]} samples · {len(st.session_state.feature_cols)} Features")
                st.dataframe(dm2[[HERKUNFT_COL] + CLASS_LEVELS].drop_duplicates(), use_container_width=True)

    # 1.2 ──────────────────────────────────────────────────────────────────────
    elif _page == EXP_PAGES[1]:
        st.title("✂️ 1.2 · Feature Selection")
        if st.session_state.df_num is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        df = st.session_state.df_num
        rem = [c for c in df.columns if c not in [HERKUNFT_COL]+CLASS_LEVELS]
        drop = st.multiselect("Remove", rem)
        if st.button("🗑️ Remove", type="primary") and drop:
            df2 = df.drop(columns=drop)
            st.session_state.df_num = df2; st.session_state.df_merged = df2
            st.session_state.feature_cols = [c for c in df2.columns
                if c not in [HERKUNFT_COL]+CLASS_LEVELS
                and np.issubdtype(df2[c].dtype, np.number)]
            for k2 in ["df_proc","df_ratios","df_grouped","lod_map","mz_state"]:
                st.session_state[k2] = None
            st.success(f"✓ {len(drop)} removed"); st.rerun()
        st.dataframe(df.head(5), use_container_width=True)

    # 1.3 ──────────────────────────────────────────────────────────────────────
    elif _page == EXP_PAGES[2]:
        st.title("⚙️ 1.3 · Preprocessing")
        if st.session_state.df_num is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        with st.expander("ℹ️ Methodik-Dokumentation (für Masterarbeit)", expanded=False):
            st.markdown("""
**Preprocessing-Pipeline (in dieser Reihenfolge):**
1. **<LOD-Behandlung:** Werte unterhalb der Nachweisgrenze (0 oder `<X`) werden durch LOD/2 (oder LOD/√2) ersetzt
2. **Imputation:** Verbleibende NaN-Werte werden durch Median (oder Mittelwert) ersetzt
3. **log₁₀-Transformation:** Reduziert Rechtsschiefe und stabilisiert Varianz
4. **Skalierung:**
   - *z-Score (Autoscaling):* x' = (x−μ)/σ — jede Variable erhält Mittelwert 0 und Varianz 1
   - *Pareto-Skalierung:* x' = (x−μ)/√σ — weniger aggressiv, erhält mehr ursprüngliche Struktur

**Literaturhinweis:** Vandeginste et al. (1998), Eriksson et al. (2006) — Standard in chemometrischen Provenance-Studien.
            """)
        c1,c2,c3 = st.columns(3)
        with c1:
            zl = st.checkbox("0 → LOD", True)
            lr = st.radio("LOD Rule", ["LOD/2","LOD/√2"])
        with c2:
            imp = st.selectbox("Imputation", ["median","mean"])
            dlog = st.checkbox("log₁₀(x)", True)
        with c3:
            eps = st.number_input("ε (log-shift)", value=1e-12, format="%.2e")
            scale = st.selectbox("Scaling", ["zscore (Autoscaling)","pareto","none"])
        scale_key = "zscore" if "zscore" in scale else ("pareto" if "pareto" in scale else "none")
        if st.button("🔧 Run Preprocessing", type="primary"):
            lr_k = "lod2" if lr == "LOD/2" else "lodsqrt2"
            dfp, lm = preprocess_rfa(st.session_state.df_num,
                                      st.session_state.feature_cols,
                                      zl, lr_k, imp, dlog, eps, scale_key)
            st.session_state.df_proc = dfp; st.session_state.lod_map = lm
            for k2 in ["df_ratios","df_grouped","mz_state"]: st.session_state[k2] = None
            st.success(f"✓ df_proc erstellt: {dfp.shape}")
        if st.session_state.df_proc is not None:
            st.dataframe(st.session_state.df_proc.head(6), use_container_width=True)
            st.download_button("⬇️ df_proc Excel", to_excel(st.session_state.df_proc),
                               "preprocessed_data.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # 1.4 Deskriptive Statistik ────────────────────────────────────────────────
    elif _page == EXP_PAGES[3]:
        st.title("📊 1.4 · Descriptive Statistics")
        with st.expander("ℹ️ Methodik: Deskriptive Statistik", expanded=False):
            st.markdown("""
**Descriptive statistics** summarize the central tendency, dispersion, and shape of data distributions.
Key measures include mean, median, standard deviation, and quartiles. **Boxplots** and **violin plots**
visualize distributions per class. The **Kruskal-Wallis test** is a non-parametric method to assess
whether samples from different groups originate from the same distribution — significant results
(p < 0.05) indicate that at least one group differs.
            """)
        if st.session_state.df_num is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        df_raw = st.session_state.df_num
        feats  = st.session_state.feature_cols
        cl     = st.session_state.class_level
        p      = cur_pal()

        src = st.selectbox("Data Source", ["df_num (raw data)","df_proc (preprocessed)"])
        d   = df_raw if "num" in src else st.session_state.df_proc
        if d is None: st.error("df_proc fehlt — 1.3 abschließen."); st.stop()

        tabs = st.tabs(["📋 Summary","📦 Box/Violin","📊 Histogram",
                        "🔥 Correlation","🗂️ Heatmap","🧪 Significance","🔀 Pairplot"])

        # Tab 0: Kennzahlen
        with tabs[0]:
            st.subheader("Summary Statistics per Feature")
            grp = st.radio("Group by", ["Overall"]+[cl], horizontal=True, key="desc_grp")
            if grp == "Overall":
                summ = d[feats].describe().T
                summ["median"] = d[feats].median()
                st.dataframe(summ.round(3), use_container_width=True)
            else:
                frames = []
                for gname, sub in d.groupby(cl):
                    s = sub[feats].describe().T[["mean","std","min","max"]].copy()
                    s.columns = [f"{c}_{gname}" for c in s.columns]
                    frames.append(s)
                st.dataframe(pd.concat(frames, axis=1).round(3), use_container_width=True)
            st.download_button("⬇️ Excel", to_excel(d[feats].describe().T.round(3)),
                               "descriptive.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # Tab 1: Box/Violin
        with tabs[1]:
            c1,c2,c3 = st.columns(3)
            with c1: plot_type = st.radio("Typ",["Boxplot","Violin Plot"], horizontal=True)
            with c2: feat_sel  = st.selectbox("Element", feats, key="bv_feat")
            with c3: show_pts  = st.checkbox("Show data points", True)
            pts = "all" if show_pts else "outliers"
            if plot_type == "Boxplot":
                fig_bv = px.box(d, x=cl, y=feat_sel, color=cl,
                                points=pts, title=f"{feat_sel} nach {cl}",
                                template="plotly_white")
            else:
                fig_bv = px.violin(d, x=cl, y=feat_sel, color=cl,
                                   box=True, points=pts,
                                   title=f"{feat_sel} nach {cl} (Violin)",
                                   template="plotly_white")
            st.plotly_chart(fig_bv, use_container_width=True)
            # Pub
            fig_m, ax = plt.subplots(figsize=(8,5))
            grps = sorted(d[cl].astype(str).unique())
            data_by_grp = [d[d[cl].astype(str)==g][feat_sel].dropna().values for g in grps]
            bp = ax.boxplot(data_by_grp, patch_artist=True, notch=False)
            for patch, g in zip(bp["boxes"], grps):
                col, _ = p.get(list(d[d[cl].astype(str)==g][HERKUNFT_COL])[0], ("#aaa","o")) \
                          if len(d[d[cl].astype(str)==g]) else ("#aaa","o")
                patch.set_facecolor(col); patch.set_alpha(0.7)
            if show_pts:
                for i, (grp_d, g) in enumerate(zip(data_by_grp, grps)):
                    ax.scatter(np.full(len(grp_d), i+1) + np.random.uniform(-0.08,0.08,len(grp_d)),
                               grp_d, s=20, alpha=0.6, color=TAB10[i%10], zorder=3)
            ax.set_xticklabels(grps, rotation=20, ha="right", fontsize=9)
            ax.set_ylabel(feat_sel, fontsize=11); ax.set_title(f"{feat_sel} nach {cl}", fontweight="bold")
            ax.grid(True, axis="y", alpha=0.2); fig_m.tight_layout()
            st.pyplot(fig_m); pub_dl(fig_m, f"box_{feat_sel}_{cl}"); plt.close(fig_m)

        # Tab 2: Histogramm
        with tabs[2]:
            c1,c2 = st.columns(2)
            with c1: h_feat = st.selectbox("Element", feats, key="h_feat")
            with c2: n_bins  = st.slider("Bins", 5, 50, 20)
            fig_h = px.histogram(d, x=h_feat, color=cl,
                                  color_discrete_map={g: list(p.values())[i][0]
                                    for i,g in enumerate(sorted(d[cl].astype(str).unique()))},
                                  nbins=n_bins, barmode="overlay", opacity=0.7,
                                  title=f"Distribution: {h_feat}", template="plotly_white")
            st.plotly_chart(fig_h, use_container_width=True)
            # With KDE
            fig_k, ax_k = plt.subplots(figsize=(8,4))
            for g in sorted(d[cl].astype(str).unique()):
                vals = d[d[cl].astype(str)==g][h_feat].dropna().values
                if len(vals) < 2: continue
                col2 = _BASE_COLORS.get(g, TAB10[0])
                ax_k.hist(vals, bins=n_bins, alpha=0.3, color=col2, density=True, label=g)
                try:
                    kde = sp_stats.gaussian_kde(vals)
                    xx = np.linspace(vals.min(), vals.max(), 200)
                    ax_k.plot(xx, kde(xx), color=col2, lw=2)
                except Exception:
                    pass
            ax_k.set_xlabel(h_feat, fontsize=11); ax_k.set_ylabel("Dichte")
            ax_k.set_title(f"Distribution + KDE: {h_feat}", fontweight="bold")
            ax_k.legend(fontsize=9); ax_k.grid(True, alpha=0.2)
            fig_k.tight_layout(); st.pyplot(fig_k)
            pub_dl(fig_k, f"hist_{h_feat}"); plt.close(fig_k)

        # Tab 3: Korrelationsmatrix
        with tabs[3]:
            top_corr = st.slider("Top Features (nach Varianz)", 5, min(38,len(feats)), min(20,len(feats)))
            top_feats = d[feats].var().nlargest(top_corr).index.tolist()
            corr_m = d[top_feats].corr(method="spearman")
            fig_co = px.imshow(corr_m, color_continuous_scale="RdBu_r",
                               color_continuous_midpoint=0, zmin=-1, zmax=1,
                               title="Spearman Correlation Matrix",
                               text_auto=".2f", aspect="equal")
            fig_co.update_traces(textfont_size=7)
            st.plotly_chart(fig_co, use_container_width=True)
            fig_cm2, ax_cm = plt.subplots(figsize=(max(7,top_corr*0.55), max(6,top_corr*0.5)))
            im = ax_cm.imshow(corr_m.values, cmap="RdBu_r", vmin=-1, vmax=1, aspect="auto")
            ax_cm.set_xticks(range(len(top_feats)))
            ax_cm.set_xticklabels(top_feats, rotation=45, ha="right", fontsize=7)
            ax_cm.set_yticks(range(len(top_feats)))
            ax_cm.set_yticklabels(top_feats, fontsize=7)
            plt.colorbar(im, ax=ax_cm, fraction=0.03); ax_cm.set_title("Spearman Correlation", fontweight="bold")
            fig_cm2.tight_layout(); st.pyplot(fig_cm2)
            pub_dl(fig_cm2, "korrelation"); plt.close(fig_cm2)

        # Tab 4: Konzentrations-Heatmap
        with tabs[4]:
            top_h = st.slider("Top Features", 5, min(38,len(feats)), min(15,len(feats)), key="hm_top")
            tv = d[feats].var().nlargest(top_h).index.tolist()
            Xhm = d[tv].astype(float).copy()
            Xhm = (Xhm - Xhm.mean()) / (Xhm.std().replace(0,1))  # z-score for display
            labels_hm = d[HERKUNFT_COL].astype(str).tolist()
            fig_hm = px.imshow(Xhm.values, x=tv, y=labels_hm,
                               color_continuous_scale="RdYlBu_r",
                               aspect="auto", title="Element Concentration Heatmap (z-normalized)")
            fig_hm.update_layout(height=max(400, len(labels_hm)*18))
            st.plotly_chart(fig_hm, use_container_width=True)
            fig_hm2, ax_hm = plt.subplots(figsize=(max(8,top_h*0.6), max(5,len(labels_hm)*0.35)))
            im2 = ax_hm.imshow(Xhm.values, cmap="RdYlBu_r", aspect="auto")
            ax_hm.set_xticks(range(len(tv))); ax_hm.set_xticklabels(tv, rotation=45, ha="right", fontsize=7)
            ax_hm.set_yticks(range(len(labels_hm))); ax_hm.set_yticklabels(labels_hm, fontsize=7)
            plt.colorbar(im2, ax=ax_hm, fraction=0.02)
            ax_hm.set_title("Concentration Heatmap", fontweight="bold")
            fig_hm2.tight_layout(); st.pyplot(fig_hm2)
            pub_dl(fig_hm2, "heatmap_konz"); plt.close(fig_hm2)

        # Tab 5: Signifikanztests
        with tabs[5]:
            st.subheader("Kruskal-Wallis Test (per element, grouped by active class level)")
            st.caption("H: Test statistic · p < 0.05 = significant · Sig: *** p<0.001, ** p<0.01, * p<0.05, ns")
            kw_df = kruskal_wallis_all(d, feats, cl)
            st.dataframe(kw_df.style.background_gradient(subset=["p_value"], cmap="RdYlGn_r"),
                         use_container_width=True)
            sig_only = kw_df[kw_df["Sig"].isin(["*","**","***"])]
            if len(sig_only):
                st.success(f"✅ {len(sig_only)} signifikante Features (p < 0.05)")
                fig_kw, ax_kw = plt.subplots(figsize=(8, max(4,len(sig_only)*0.4+1)))
                ax_kw.barh(sig_only["Feature"][:20], -np.log10(sig_only["p_value"][:20]+1e-12),
                           color=[TAB10[0] if s=="*" else (TAB10[1] if s=="**" else TAB10[2])
                                  for s in sig_only["Sig"][:20]])
                ax_kw.axvline(-np.log10(0.05), color="red", ls="--", lw=1, label="p=0.05")
                ax_kw.set_xlabel("-log₁₀(p-value)"); ax_kw.set_title("Kruskal-Wallis Significance", fontweight="bold")
                ax_kw.legend(fontsize=8); ax_kw.grid(True, axis="x", alpha=0.2)
                fig_kw.tight_layout(); st.pyplot(fig_kw)
                pub_dl(fig_kw, "kruskal_wallis"); plt.close(fig_kw)
            st.download_button("⬇️ Test Results Excel", to_excel(kw_df),
                               "kruskal_wallis.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # Tab 6: Pairplot
        with tabs[6]:
            st.subheader("Pairplot of Selected Features")
            default_pp = ["Rb","K","Cs","Fe","Mg","Ti"]
            avail_pp = [f for f in default_pp if f in feats] + [f for f in feats if f not in default_pp]
            sel_pp = st.multiselect("Features (4–8 empfohlen)", feats, default=avail_pp[:min(6,len(avail_pp))])
            if len(sel_pp) < 2: st.warning("Min. 2 Features."); st.stop()
            if st.button("🔀 Pairplot erstellen", type="primary"):
                df_pp = d[sel_pp + [HERKUNFT_COL, cl]].copy()
                fig_pp = px.scatter_matrix(df_pp, dimensions=sel_pp,
                                            color=cl, symbol=HERKUNFT_COL,
                                            color_discrete_map={
                                                g: _BASE_COLORS.get(g,"#777")
                                                for g in df_pp[cl].unique()},
                                            title=f"Pairplot [{cl}]",
                                            opacity=0.8, height=700)
                fig_pp.update_traces(marker_size=6)
                st.plotly_chart(fig_pp, use_container_width=True)
                n_pp = len(sel_pp)
                fig_pm, axes = plt.subplots(n_pp, n_pp, figsize=(max(8,n_pp*2), max(8,n_pp*2)))
                for i, fi in enumerate(sel_pp):
                    for j, fj in enumerate(sel_pp):
                        ax_ij = axes[i,j] if n_pp > 1 else axes
                        if i == j:
                            for g in sorted(d[cl].astype(str).unique()):
                                vals = d[d[cl].astype(str)==g][fi].dropna().values
                                if len(vals) > 1:
                                    ax_ij.hist(vals, bins=10, alpha=0.5,
                                               color=_BASE_COLORS.get(g,"#777"), density=True)
                        else:
                            done_pp = set()
                            for _, row_pp in d.iterrows():
                                h_pp = str(row_pp[HERKUNFT_COL]); col_pp, mkr_pp = p.get(h_pp,("#aaa","o"))
                                lbl_pp = h_pp if h_pp not in done_pp else "_nl_"; done_pp.add(h_pp)
                                ax_ij.scatter(row_pp[fj], row_pp[fi], c=col_pp, marker=mkr_pp,
                                              s=18, alpha=0.8, linewidths=0.2, edgecolors="white")
                        if i == n_pp-1: ax_ij.set_xlabel(fj, fontsize=7)
                        if j == 0: ax_ij.set_ylabel(fi, fontsize=7)
                        ax_ij.tick_params(labelsize=6)
                fig_pm.suptitle(f"Pairplot [{cl}]", fontsize=11, fontweight="bold")
                fig_pm.tight_layout(); st.pyplot(fig_pm)
                pub_dl(fig_pm, "pairplot"); plt.close(fig_pm)

    # 1.5 Color Scheme ──────────────────────────────────────────────────────────
    elif _page == EXP_PAGES[4]:
        st.title("🎨 1.5 · Color Scheme Preview")
        if st.session_state.df_merged is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        dm = st.session_state.df_merged; cl = st.session_state.class_level; p = cur_pal()
        st.info(f"Klassen-Ebene: **{cl}** · Symbol-Detail: **{'an' if st.session_state.symbol_detail else 'aus'}**\n\n"
                "Einstellungen in der Sidebar ändern.")
        fig_pv, ax_pv = plt.subplots(figsize=(10, max(3, len(p)*0.38+1)))
        for i, (h, (col, mkr)) in enumerate(sorted(p.items())):
            row = dm[dm[HERKUNFT_COL]==h]; cls2 = str(row.iloc[0][cl]) if len(row) else "?"
            ax_pv.scatter(0.5, i, color=col, marker=mkr, s=100, lw=0.4, edgecolors="white")
            ax_pv.text(0.65, i, f"{h}  [{cls2}]", va="center", fontsize=9)
        ax_pv.set_xlim(0,4); ax_pv.set_yticks([]); ax_pv.set_xticks([])
        ax_pv.set_title("Color & Symbol per Sample Origin", fontweight="bold"); ax_pv.set_frame_on(False)
        fig_pv.tight_layout(); st.pyplot(fig_pv); plt.close(fig_pv)
        rows_t = []
        for h, (col, mkr) in sorted(p.items()):
            row = dm[dm[HERKUNFT_COL]==h]; cls2 = str(row.iloc[0][cl]) if len(row) else "?"
            rows_t.append({cl: cls2, HERKUNFT_COL: h, "Farbe": col, "Symbol": mkr})
        st.dataframe(pd.DataFrame(rows_t), use_container_width=True)

    # 1.6 Screening ───────────────────────────────────────────────────────────
    elif _page == EXP_PAGES[5]:
        st.title("🔬 1.6 · Screening-Plot")
        if st.session_state.df_num is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        feats = st.session_state.feature_cols; p = cur_pal(); cl = st.session_state.class_level
        c1,c2,c3 = st.columns(3)
        with c1: src = st.selectbox("Data",["df_num (raw)","df_proc (processed)"])
        with c2: xcol = st.selectbox("X", feats, index=feats.index("Mg") if "Mg" in feats else 0)
        with c3: ycol = st.selectbox("Y", feats, index=feats.index("Al") if "Al" in feats else min(1,len(feats)-1))
        d = st.session_state.df_num if "num" in src else st.session_state.df_proc
        if d is None: st.error("df_proc fehlt → 1.3 abschließen.")
        elif st.button("📊 Plot", type="primary"):
            title = f"Screening: {xcol} vs {ycol}  [{cl}]"
            st.subheader("🖱️ Interactive")
            fig_sc = px.scatter(d, x=xcol, y=ycol, color=HERKUNFT_COL, symbol=HERKUNFT_COL,
                                color_discrete_map=pal_color_map(p),
                                symbol_map=pal_symbol_map(p),
                                title=title, template="plotly_white")
            fig_sc.update_traces(marker_size=10, marker_opacity=0.88)
            st.plotly_chart(fig_sc, use_container_width=True)
            st.divider(); st.subheader("🖨️ Publication")
            fm = scatter_mpl(d, xcol, ycol, p, cl, title)
            st.pyplot(fm); pub_dl(fm, f"screening_{xcol}_{ycol}"); plt.close(fm)

    # 1.7 Ratios ──────────────────────────────────────────────────────────────
    elif _page == EXP_PAGES[6]:
        st.title("➗ 1.7 · Ratio-Features")
        if st.session_state.df_num is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        feats = st.session_state.feature_cols; p = cur_pal(); cl = st.session_state.class_level
        st.warning("⚠️ Ratios are calculated on **raw data (df_num)** or **log₁₀-transformed data** only "
                   "berechnet — nie auf z-skalierten Werten, da diese geochemisch nicht interpretierbar sind.")
        src = st.selectbox("Data Source", ["df_num (raw — recommended)","df_proc (only if NO z-scaling)"])
        d = st.session_state.df_num
        if "proc" in src and st.session_state.df_proc is not None:
            d = st.session_state.df_proc
        min_d = st.number_input("Minimaler Nenner (unter diesem Wert → NaN)", value=0.01, format="%.4f")
        PRED = [("Rb_over_K","Rb","K"),("Cs_over_K","Cs","K"),("Ba_over_Rb","Ba","Rb"),
                ("Fe_over_Mg","Fe","Mg"),("Ti_over_Al","Ti","Al")]
        for n_,a_,b_ in PRED:
            av = a_ in d.columns and b_ in d.columns
            st.markdown(f"{'✅' if av else '❌'} `{n_}` = {a_}/{b_}")
        c1,c2,c3 = st.columns(3)
        with c1: cn = st.text_input("Eigener Name","custom_ratio")
        with c2: cnum = st.selectbox("Zähler", feats, key="rn")
        with c3: cden = st.selectbox("Nenner", feats, key="rd")
        if st.button("⚙️ Compute", type="primary"):
            ratios = {}
            for n_,a_,b_ in PRED:
                r = safe_ratio(d, a_, b_, min_denom=min_d)
                if r is not None: ratios[n_] = r
            if cn and cnum != cden:
                r = safe_ratio(d, cnum, cden, min_denom=min_d)
                if r is not None: ratios[cn] = r
            if ratios:
                ec = [c for c in [HERKUNFT_COL]+CLASS_LEVELS if c in d.columns]
                df_r = pd.concat([d[ec].copy(), pd.DataFrame(ratios)], axis=1)
                st.session_state.df_ratios = df_r
                nan_pct = {k: v.isna().mean()*100 for k,v in ratios.items()}
                st.success(f"✓ {len(ratios)} Ratio(s) berechnet")
                for k,v in nan_pct.items():
                    if v > 20: st.warning(f"⚠️ {k}: {v:.1f}% NaN (Nenner zu klein)")
                st.dataframe(df_r.head(8), use_container_width=True)
        if st.session_state.df_ratios is not None:
            st.download_button("⬇️ Ratios Excel", to_excel(st.session_state.df_ratios),
                               "element_ratios.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    # 1.8 PCA Global ──────────────────────────────────────────────────────────
    elif _page == EXP_PAGES[7]:
        st.title("📈 1.8 · PCA Global")
        with st.expander("ℹ️ Methodik: PCA", expanded=False):
            st.markdown("""
**Principal Component Analysis (PCA)** is an unsupervised dimensionality reduction technique that
projects high-dimensional data onto orthogonal components (PCs) ordered by decreasing variance.
**Scores** show sample positions in the reduced space; **loadings** indicate the contribution of
original variables to each PC. PCA is commonly used as an exploratory first step before
supervised classification in chemometric provenance studies.
            """)
        if st.session_state.df_proc is None: st.warning("⚠️ Complete step 1.3 first."); st.stop()
        feats = st.session_state.feature_cols; p = cur_pal(); cl = st.session_state.class_level
        d = st.session_state.df_proc
        c1,c2 = st.columns(2)
        with c1: top_n = st.slider("Loading Arrows",3,20,8); scale_bp = st.slider("Arrow Scaling",0.5,8.0,3.0,0.5)
        with c2: top_ld = st.slider("Top Loadings Balken",5,30,12)
        if st.button("🔍 Compute PCA", type="primary"):
            X = SimpleImputer(strategy="median").fit_transform(d[feats].astype(float).values)
            Z,ve,ld_df,_ = run_pca(X, feats)
            ld2 = ld_df[["PC1","PC2"]].copy(); ld2["Feature"] = ld_df.index
            ld2["max_abs"] = ld2[["PC1","PC2"]].abs().max(axis=1)
            st.session_state.pca_scores   = pd.DataFrame({"PC1":Z[:,0],"PC2":Z[:,1]})
            st.session_state.pca_loadings = ld2
            sc_df = pd.concat([pd.DataFrame({"PC1":Z[:,0],"PC2":Z[:,1]}),
                                d[[HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)], axis=1)
            st.subheader("🖱️ Interactive")
            top_f = ld_df["PC1"].abs().add(ld_df["PC2"].abs()).nlargest(top_n).index
            fig_px = px.scatter(sc_df, x="PC1", y="PC2", color=HERKUNFT_COL, symbol=HERKUNFT_COL,
                                color_discrete_map=pal_color_map(p), symbol_map=pal_symbol_map(p),
                                hover_data=CLASS_LEVELS, title=f"PCA Global [{cl}]", template="plotly_white",
                                labels={"PC1":f"PC1({ve[0]*100:.1f}%)","PC2":f"PC2({ve[1]*100:.1f}%)"})
            for feat in top_f:
                lx = ld_df.loc[feat,"PC1"]*scale_bp; ly = ld_df.loc[feat,"PC2"]*scale_bp
                fig_px.add_annotation(x=lx, y=ly, ax=0, ay=0, xref="x", yref="y", axref="x", ayref="y",
                                      showarrow=True, arrowhead=2, arrowwidth=1.5, arrowcolor="#555")
                fig_px.add_annotation(x=lx*1.1, y=ly*1.1, text=feat, showarrow=False,
                                      font=dict(size=10, color="#333"))
            fig_px.update_traces(marker_size=10, marker_opacity=0.88)
            st.plotly_chart(fig_px, use_container_width=True)
            st.divider(); st.subheader("🖨️ Biplot")
            fig_bp = biplot_mpl(Z, ld_df, d.reset_index(drop=True), p, cl,
                                 f"PCA Biplot [{cl}]", ve, top_n, scale_bp)
            st.pyplot(fig_bp); pub_dl(fig_bp, "pca_biplot"); plt.close(fig_bp)
            st.subheader("🖨️ Loadings")
            fig_ld = loadings_mpl(ld2, top_ld, "Top Loadings")
            st.pyplot(fig_ld); pub_dl(fig_ld, "pca_loadings"); plt.close(fig_ld)
            c1,c2 = st.columns(2); c1.metric("PC1",f"{ve[0]*100:.1f}%"); c2.metric("PC2",f"{ve[1]*100:.1f}%")
            st.dataframe(ld2.sort_values("max_abs",ascending=False).head(top_ld).reset_index(drop=True),
                         use_container_width=True)

    # 1.9 Sample Groups ──────────────────────────────────────────────────────
    elif _page == EXP_PAGES[8]:
        st.title("🧩 1.9 · Sample Groups & PCA")
        if st.session_state.df_num is None: st.warning("⚠️ Complete step 1.1 first."); st.stop()
        feats = st.session_state.feature_cols; p = cur_pal(); cl = st.session_state.class_level
        st.success("✅ Grouping is performed on **raw data (df_num)** — "
                   "the threshold is applied to a selected element for physically interpretable grouping.")
        d_raw = st.session_state.df_num
        c1,c2 = st.columns(2)
        with c1: mgo = st.number_input("Mg ≤ (Group A threshold)", value=50.0, step=5.0)
        with c2: top_n = st.slider("Top Loadings",5,25,10)
        if "Mg" not in d_raw.columns:
            st.error("Column 'Mg' not found in dataset."); st.stop()
        if st.button("🔬 PCA per Group", type="primary"):
            dc = d_raw.copy()
            dc["_MG"] = np.where(dc["Mg"].astype(float) <= mgo, "Group_A","Group_B")
            st.session_state.df_grouped = dc
            grp_sc, grp_ld = {}, {}
            for gname, sub in dc.groupby("_MG"):
                st.markdown(f"---\n### {gname} (n={len(sub)})")
                if len(sub) < 3: st.warning("Too few samples."); continue
                X = SimpleImputer(strategy="median").fit_transform(sub[feats].astype(float).values)
                if st.session_state.df_proc is not None:
                    idx_in_proc = sub.index
                    sub_proc = st.session_state.df_proc.loc[
                        st.session_state.df_proc.index.isin(idx_in_proc)] \
                        if any(st.session_state.df_proc.index.isin(idx_in_proc)) \
                        else None
                    if sub_proc is not None and len(sub_proc) == len(sub):
                        X = sub_proc[feats].astype(float).fillna(0).values
                Z,ve,ld_df,_ = run_pca(X, feats)
                ld2 = ld_df[["PC1","PC2"]].copy(); ld2["Feature"] = ld_df.index
                ld2["max_abs"] = ld2[["PC1","PC2"]].abs().max(axis=1)
                grp_sc[str(gname)] = pd.DataFrame({"PC1":Z[:,0],"PC2":Z[:,1]})
                grp_ld[str(gname)] = ld2; sg = str(gname).replace(" ","_")
                sc_df = pd.concat([pd.DataFrame({"PC1":Z[:,0],"PC2":Z[:,1]}),
                                    sub[[HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)], axis=1)
                st.subheader(f"🖱️ {gname}")
                fp = px.scatter(sc_df, x="PC1", y="PC2", color=HERKUNFT_COL, symbol=HERKUNFT_COL,
                                color_discrete_map=pal_color_map(p), symbol_map=pal_symbol_map(p),
                                title=f"PCA · {gname}", template="plotly_white")
                fp.update_traces(marker_size=9, marker_opacity=0.88); st.plotly_chart(fp,use_container_width=True)
                st.subheader(f"🖨️ {gname}")
                fsf = pca_mpl(Z, sub.reset_index(drop=True), p, cl, f"PCA · {gname}", ve)
                st.pyplot(fsf); pub_dl(fsf, f"pca_{sg}"); plt.close(fsf)
                fl = loadings_mpl(ld2, top_n, f"Loadings · {gname}")
                st.pyplot(fl); pub_dl(fl, f"ld_{sg}"); plt.close(fl)
            st.session_state.pca_group_scores  = grp_sc
            st.session_state.pca_group_loadings = grp_ld

    # 1.10 Univariate Scatter ─────────────────────────────────────────────────
    elif _page == EXP_PAGES[9]:
        st.title("📉 1.10 · Univariate Scatter")
        if st.session_state.df_grouped is None: st.warning("⚠️ Complete step 1.9 first."); st.stop()
        feats = st.session_state.feature_cols; p = cur_pal(); cl = st.session_state.class_level
        d = st.session_state.df_grouped; groups = list(d["_MG"].unique())
        c1,c2,c3,c4 = st.columns(4)
        with c1: gsel = st.selectbox("Gruppe", groups)
        with c2: src2 = st.selectbox("Daten",["df_proc","df_num"])
        with c3: xcol = st.selectbox("X", feats, index=feats.index("Rb") if "Rb" in feats else 0)
        with c4: ycol = st.selectbox("Y", feats, index=feats.index("Cs") if "Cs" in feats else min(1,len(feats)-1))
        if st.button("📈 Plot", type="primary"):
            ds = st.session_state.df_proc if src2=="df_proc" else st.session_state.df_num
            if ds is None: st.error("Data source missing.")
            else:
                ds2 = ds.copy(); ds2["_MG"] = d["_MG"].values
                sub = ds2[ds2["_MG"].astype(str)==str(gsel)]
                if len(sub) < 2: st.error("Zu wenig.")
                else:
                    title = f"{gsel}: {xcol} vs {ycol}  [{cl}]"
                    st.subheader("🖱️ Interactive")
                    fp = px.scatter(sub, x=xcol, y=ycol, color=HERKUNFT_COL, symbol=HERKUNFT_COL,
                                    color_discrete_map=pal_color_map(p), symbol_map=pal_symbol_map(p),
                                    title=title, template="plotly_white")
                    fp.update_traces(marker_size=10, marker_opacity=0.88); st.plotly_chart(fp,use_container_width=True)
                    st.divider(); st.subheader("🖨️ Publication")
                    fm = scatter_mpl(sub, xcol, ycol, p, cl, title)
                    st.pyplot(fm); pub_dl(fm, f"univ_{gsel.replace(' ','_')}_{xcol}_{ycol}"); plt.close(fm)

# ══════════════════════════════════════════════════════════════════════════════
# SEITE 2 — MODEL ZOO
# ══════════════════════════════════════════════════════════════════════════════
elif _section == SEC_MZ:
    df_num = st.session_state.df_num; lod_map = st.session_state.lod_map
    feats  = st.session_state.feature_cols; cl = st.session_state.class_level; p = cur_pal()

    # ══════ SUPERVISED ══════════════════════════════════════════════════════
    if "Supervised" in _mz_mode:

        # S.1 Training ───────────────────────────────────────────────────────
        if _page == SUP_PAGES[0]:
            st.title("⚙️ S.1 · Training & k-Fold CV")
            if df_num is None: st.error("❌ Seite 1 abschließen."); st.stop()
            # Class size check
            counts, small = check_class_sizes(df_num, cl, min_samples=3)
            st.markdown(f"**Active class level:** `{cl}` — {df_num[cl].nunique()} classes")
            st.dataframe(counts.reset_index().rename(columns={"index":cl, cl:"Count"}), use_container_width=True)
            if len(small):
                st.warning(f"⚠️ Small classes (< 3 samples): {small.to_dict()} — Train/test split may be unstable!")

            # Bauteile
            st.divider(); st.subheader("🔩 Unknown Samples (optional)")
            up_unk = st.file_uploader("Unknown Samples File", type=["xlsx","xls","csv","txt"], key="unk3")
            if up_unk:
                try:
                    ext2 = up_unk.name.split(".")[-1].lower()
                    if ext2 in ("xlsx","xls"):
                        xl2 = pd.ExcelFile(up_unk); sh2 = st.selectbox("Sheet",xl2.sheet_names,key="ush2")
                        df_unk = pd.read_excel(xl2, sheet_name=sh2)
                    else:
                        s2 = st.selectbox("Sep",[",",";","\t"],key="usp2"); df_unk = pd.read_csv(up_unk,sep=s2)
                    no = ["— automatic"]+list(df_unk.columns); ns = st.selectbox("Name Column",no)
                    nc = None if ns.startswith("—") else ns
                    for c2 in [col for col in df_unk.columns if col!=nc]: df_unk[c2] = coerce_num(df_unk[c2])
                    if st.button("✅ Apply Unknown Samples"):
                        st.session_state.mz_df_unknown = df_unk; st.session_state.mz_unk_name_col = nc
                        st.session_state.mz_state = None; st.success("✅ Unknown samples saved."); st.rerun()
                except Exception as e: st.error(f"Error: {e}")
            elif st.session_state.mz_df_unknown is not None:
                st.success(f"✅ {st.session_state.mz_df_unknown.shape[0]} unknown samples loaded")

            st.divider()
            c1,c2,c3 = st.columns(3)
            with c1:
                ts = st.slider("Test-Anteil", 0.10, 0.50, 0.30, 0.05)
                kfold_k = st.slider("k-Fold k", 3, 10, 5)
                sd = st.number_input("Seed", value=42, step=1)
            with c2:
                kk = st.slider("kNN k", 1, 25, 7)
                sc_svm = st.select_slider("SVM C", [0.1,0.5,1,5,10,50,100,500], value=10)
            with c3:
                rf = st.slider("RF Bäume", 100, 2000, 500, 100)
                scale_m = st.selectbox("Scaling", ["zscore","pareto"])
            with st.expander("⚙️ Preprocessing"):
                cp1,cp2 = st.columns(2)
                with cp1:
                    dl = st.checkbox("0→LOD", value=(lod_map is not None), key="mz_dl3")
                    lr = st.radio("LOD Rule", ["LOD/2","LOD/√2"], key="mz_lr3")
                with cp2:
                    dlog = st.checkbox("log₁₀", True, key="mz_dlog3")
                    ep = st.number_input("ε", 1e-12, format="%.2e", key="mz_ep3")
            fm_sel = st.radio("Features", ["Alle","Manuell"], horizontal=True)
            sf = st.multiselect("Features", feats, default=feats) if fm_sel=="Manuell" else feats

            if st.button("🚀 Train", type="primary"):
                if len(sf) < 2: st.error("Min. 2."); st.stop()
                with st.spinner("Training + k-Fold CV …"):
                    try:
                        al = lod_map if dl else None
                        lrk = "lod2" if lr=="LOD/2" else "lodsqrt2"
                        st_d = mz_evaluate(df_num, cl, sf, ts, sd, kk, sc_svm, rf,
                                           al, lrk, dlog, float(ep), kfold_k, scale_m)
                        st.session_state.mz_state = st_d
                        st.success("✅ Training + CV fertig!")
                        st.subheader(f"📊 k-Fold CV (k={kfold_k}) — Primäre Metrik")
                        st.dataframe(st_d["cv_results"].style.background_gradient(
                            subset=[c for c in st_d["cv_results"].columns if "F1_mean" in c],
                            cmap="RdYlGn"), use_container_width=True)
                        st.subheader("📊 Train/Test Split — Sekundäre Metrik")
                        st.dataframe(st_d["split_results"].style.background_gradient(
                            subset=["Split_F1"], cmap="RdYlGn"), use_container_width=True)
                    except Exception as e:
                        st.error(f"Error: {e}")
            elif st.session_state.mz_state is not None:
                mz = st.session_state.mz_state
                st.subheader(f"📊 k-Fold CV (k={mz.get('kfold_k',5)})")
                st.dataframe(mz["cv_results"], use_container_width=True)

        # S.2 PLS-DA ─────────────────────────────────────────────────────────
        elif _page == SUP_PAGES[1]:
            st.title("🔬 S.2 · PLS-DA + VIP Scores")
            with st.expander("ℹ️ Methodik: PLS-DA", expanded=False):
                st.markdown("""
**Partial Least Squares Discriminant Analysis (PLS-DA)** is a supervised classification method
that combines the dimensionality reduction of PLS regression with class discrimination.
It identifies latent variables (LVs) in the predictor matrix **X** that maximize the covariance
with a dummy-coded class matrix **Y**. The **VIP score** (Variable Importance in Projection)
quantifies the contribution of each original variable to the discrimination, where VIP > 1
indicates high relevance.
                """)
            mz = st.session_state.mz_state
            if mz is None: st.warning("⚠️ Complete step S.1 first."); st.stop()
            features = mz["features"]; label_col = mz["label_col"]
            X_all = mz_lod_log(mz["df_train"], features, mz.get("lod_map"),
                                mz.get("lod_rule","lod2"), mz.get("do_log10",True), mz.get("eps",1e-12))
            y_all = mz["df_train"][label_col].astype(str).values
            Xtr   = mz["X_train"].values if hasattr(mz["X_train"],"values") else np.array(mz["X_train"])
            ytr   = mz["y_train"].astype(str).values
            c1,c2,c3 = st.columns(3)
            with c1: n_comp = st.slider("PLS Komponenten", 1, min(8,len(set(ytr))+2), 2)
            with c2: vip_thr = st.slider("VIP Threshold", 0.5, 1.5, 1.0, 0.1)
            with c3: cv_k_pls = st.slider("k-Fold für CV", 3, 10, 5)

            if st.button("🔬 Compute PLS-DA", type="primary"):
                with st.spinner("PLS-DA …"):
                    prep_pls = Pipeline([("imp",SimpleImputer(strategy="median")),
                                         ("sc",StandardScaler())])
                    X_sc = prep_pls.fit_transform(Xtr)
                    X_te_sc = prep_pls.transform(mz["X_test"].values if hasattr(mz["X_test"],"values") else np.array(mz["X_test"]))
                    yte = mz["y_test"].astype(str).values
                    pls, classes_pls = fit_plsda(X_sc, ytr, n_components=n_comp)
                    vip = compute_vip(pls)
                    vip_df = pd.DataFrame({"Feature":features,"VIP":vip})
                    vip_df = vip_df.sort_values("VIP",ascending=False).reset_index(drop=True)
                    # Scores plot
                    Z_tr = pls.transform(X_sc); Z_te = pls.transform(X_te_sc)
                    # Prediction
                    y_pred_tr = predict_plsda(pls, classes_pls, X_sc)
                    y_pred_te = predict_plsda(pls, classes_pls, X_te_sc)
                    acc_te  = accuracy_score(yte, y_pred_te)
                    f1_te   = f1_score(yte, y_pred_te, average="macro", zero_division=0)
                    # k-fold
                    X_all_sc = prep_pls.fit_transform(X_all.values if hasattr(X_all,"values") else np.array(X_all))
                    cv_accs, cv_f1s = plsda_kfold(X_all_sc, y_all, n_comp, cv_k_pls, mz["seed"])
                    # Store
                    st.session_state.mz_state["plsda"] = {
                        "pls": pls, "classes": classes_pls, "vip_df": vip_df,
                        "Z_tr": Z_tr, "Z_te": Z_te, "ytr": ytr, "yte": yte,
                        "n_comp": n_comp}

                    c1m,c2m,c3m,c4m = st.columns(4)
                    c1m.metric("Test Accuracy", f"{acc_te:.3f}")
                    c2m.metric("Test MacroF1",  f"{f1_te:.3f}")
                    c3m.metric(f"CV{cv_k_pls} Acc",  f"{cv_accs.mean():.3f}±{cv_accs.std():.3f}")
                    c4m.metric(f"CV{cv_k_pls} F1",   f"{cv_f1s.mean():.3f}±{cv_f1s.std():.3f}")

                    # Scores plot (interaktiv) — robust for 1 or 2+ components
                    st.subheader("🖱️ PLS-DA Scores")
                    df_tr_ref = mz["df_train"].loc[mz["X_train"].index, [HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)
                    df_te_ref = mz["df_train"].loc[mz["X_test"].index,  [HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)
                    _n_lv = Z_tr.shape[1]  # actual number of latent variables returned

                    if _n_lv >= 2:
                        sc_all = pd.concat([
                            pd.concat([pd.DataFrame({"LV1":Z_tr[:,0],"LV2":Z_tr[:,1],"Set":["Train"]*len(Z_tr)}),df_tr_ref],axis=1),
                            pd.concat([pd.DataFrame({"LV1":Z_te[:,0],"LV2":Z_te[:,1],"Set":["Test"]*len(Z_te)}),df_te_ref],axis=1)],
                            ignore_index=True)
                        fig_pls = px.scatter(sc_all, x="LV1", y="LV2", color=HERKUNFT_COL,
                                             symbol="Set", color_discrete_map=pal_color_map(p),
                                             title=f"PLS-DA Scores [{cl}] — n_comp={n_comp}",
                                             template="plotly_white")
                        fig_pls.update_traces(marker_size=10, marker_opacity=0.88)
                        st.plotly_chart(fig_pls, use_container_width=True)
                    else:
                        # Only 1 latent variable — show 1D separation
                        st.info("ℹ️ **Nur eine latente Variable extrahiert.** "
                                "Die Trennung ist eindimensional — LV1 wird als Stripplot dargestellt.")
                        sc_all = pd.concat([
                            pd.concat([pd.DataFrame({"LV1":Z_tr[:,0],"Set":["Train"]*len(Z_tr)}),df_tr_ref],axis=1),
                            pd.concat([pd.DataFrame({"LV1":Z_te[:,0],"Set":["Test"]*len(Z_te)}),df_te_ref],axis=1)],
                            ignore_index=True)
                        fig_pls = px.strip(sc_all, x=cl, y="LV1", color=HERKUNFT_COL,
                                           symbol="Set", color_discrete_map=pal_color_map(p),
                                           title=f"PLS-DA Scores (1D) [{cl}] — n_comp=1",
                                           template="plotly_white")
                        fig_pls.update_traces(marker_size=10, marker_opacity=0.88)
                        st.plotly_chart(fig_pls, use_container_width=True)

                    # VIP plot
                    st.subheader("📊 VIP Scores")
                    n_vip_show = st.slider("Top VIP zeigen", 5, min(30,len(features)), 15)
                    vip_top = vip_df.head(n_vip_show)
                    fig_vip = px.bar(vip_top, x="VIP", y="Feature", orientation="h",
                                     color="VIP", color_continuous_scale="RdYlGn",
                                     color_continuous_midpoint=vip_thr,
                                     title=f"VIP Scores (Schwelle={vip_thr})",
                                     template="plotly_white")
                    fig_vip.add_vline(x=vip_thr, line_dash="dash", line_color="red",
                                      annotation_text=f"VIP={vip_thr}")
                    st.plotly_chart(fig_vip, use_container_width=True)
                    sig_vip = vip_df[vip_df["VIP"] >= vip_thr]
                    st.success(f"✅ {len(sig_vip)} Features mit VIP ≥ {vip_thr}: {', '.join(sig_vip['Feature'].tolist()[:10])}")

                    fig_vm, ax_vm = plt.subplots(figsize=(9,4))
                    colors_vip = [TAB10[0] if v >= vip_thr else TAB10[1] for v in vip_top["VIP"]]
                    ax_vm.barh(vip_top["Feature"][::-1], vip_top["VIP"][::-1], color=colors_vip[::-1], alpha=0.85)
                    ax_vm.axvline(vip_thr, color="red", ls="--", lw=1.5, label=f"VIP={vip_thr}")
                    ax_vm.set_xlabel("VIP Score", fontsize=11)
                    ax_vm.set_title("PLS-DA VIP Scores", fontsize=12, fontweight="bold")
                    ax_vm.legend(fontsize=9); ax_vm.grid(True, axis="x", alpha=0.2)
                    fig_vm.tight_layout(); st.pyplot(fig_vm)
                    pub_dl(fig_vm, "plsda_vip"); plt.close(fig_vm)
                    st.download_button("⬇️ VIP Excel", to_excel(vip_df),
                                       "plsda_vip.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

            elif "plsda" in (st.session_state.mz_state or {}):
                plsd = st.session_state.mz_state["plsda"]
                st.info(f"PLS-DA bereits berechnet: {plsd['n_comp']} Komponenten, {len(plsd['classes'])} Klassen")
                st.dataframe(plsd["vip_df"].head(15), use_container_width=True)

        # S.3 Ergebnisse ─────────────────────────────────────────────────────
        elif _page == SUP_PAGES[2]:
            st.title("📊 S.3 · Ergebnisse & Feature Importance")
            mz = st.session_state.mz_state
            if mz is None: st.warning("⚠️ Complete step S.1 first."); st.stop()

            st.subheader(f"k-Fold CV (k={mz.get('kfold_k',5)}) — Primär")
            st.dataframe(mz["cv_results"].style.background_gradient(
                subset=[c for c in mz["cv_results"].columns if "F1_mean" in c], cmap="RdYlGn"),
                use_container_width=True)
            st.subheader("Train/Test Split — Sekundär")
            st.dataframe(mz["split_results"].style.background_gradient(
                subset=["Split_F1"], cmap="RdYlGn"), use_container_width=True)

            st.divider()
            ms = st.selectbox("Modell (für Konfusion + Importance)", list(mz["pipes"].keys()))
            lbls, cm_abs = mz["confmats"][ms]

            # Confusion matrix toggle
            norm_cm = st.checkbox("Zeilen-normalisierte Konfusionsmatrix", True)
            if norm_cm:
                row_sums = cm_abs.sum(axis=1, keepdims=True)
                cm_show = np.where(row_sums > 0, cm_abs / row_sums, 0)
                cm_text = [[f"{v:.2f}" for v in row] for row in cm_show]
                cm_title = f"Norm. Konfusion — {ms}"
                cscale = "Blues"
            else:
                cm_show = cm_abs; cm_text = cm_abs; cm_title = f"Konfusion — {ms}"; cscale = "Blues"
            fig_cm = px.imshow(cm_show, x=[f"P:{l}" for l in lbls], y=[f"T:{l}" for l in lbls],
                               text_auto=True, color_continuous_scale=cscale, title=cm_title)
            st.plotly_chart(fig_cm, use_container_width=True)

            # Pub: both side by side
            fig_mco, (ax1,ax2) = plt.subplots(1,2,figsize=(max(8,len(lbls)*1.4+2),max(4,len(lbls)*0.9+1)))
            for ax_i, (cm_i, ttl_i) in zip([ax1,ax2],
                [(cm_abs,"Absolut"),(cm_show if norm_cm else cm_abs, "Normiert")]):
                im_i = ax_i.imshow(cm_i, cmap="Blues", vmin=0, vmax=1 if norm_cm and ax_i==ax2 else None)
                ax_i.set_xticks(range(len(lbls))); ax_i.set_xticklabels([f"P:{l}" for l in lbls],rotation=45,ha="right",fontsize=8)
                ax_i.set_yticks(range(len(lbls))); ax_i.set_yticklabels([f"T:{l}" for l in lbls],fontsize=8)
                for ii in range(len(lbls)):
                    for jj in range(len(lbls)):
                        v_show = f"{cm_i[ii,jj]:.2f}" if (norm_cm and ax_i==ax2) else str(cm_i[ii,jj])
                        ax_i.text(jj,ii,v_show,ha="center",va="center",
                                  color="white" if cm_i[ii,jj]>cm_i.max()*0.6 else "black",fontsize=9)
                ax_i.set_title(f"{ms} — {ttl_i}", fontweight="bold"); plt.colorbar(im_i,ax=ax_i,fraction=0.046,pad=0.04)
            fig_mco.tight_layout(); st.pyplot(fig_mco)
            pub_dl(fig_mco, f"confmat_{ms.split('(')[0].strip()}"); plt.close(fig_mco)

            # Classification report
            yt = mz["y_test"].astype(str).values; yp = mz["preds_test"][ms]
            rep = classification_report(yt, yp, output_dict=True, zero_division=0)
            st.dataframe(pd.DataFrame(rep).T.round(3), use_container_width=True)

            # Feature Importance
            st.divider(); st.subheader("🔍 Feature Importance")
            fi_df = mz_feature_importance(mz, ms)
            if fi_df is not None:
                n_fi = st.slider("Top Features anzeigen", 5, min(30,len(mz["features"])), 15, key="fi_n")
                fig_fi = px.bar(fi_df.head(n_fi), x="Importance", y="Feature",
                                orientation="h", color="Importance",
                                color_continuous_scale="Oranges",
                                title=f"Feature Importance — {ms}", template="plotly_white")
                st.plotly_chart(fig_fi, use_container_width=True)
                fig_fim, ax_fi = plt.subplots(figsize=(8,max(4,n_fi*0.4)))
                top_fi = fi_df.head(n_fi)
                ax_fi.barh(top_fi["Feature"][::-1], top_fi["Importance"][::-1],
                           color=TAB10[1], alpha=0.85)
                ax_fi.set_xlabel("Importance", fontsize=11)
                ax_fi.set_title(f"Feature Importance — {ms}", fontweight="bold")
                ax_fi.grid(True, axis="x", alpha=0.2); fig_fim.tight_layout()
                st.pyplot(fig_fim); pub_dl(fig_fim, f"fi_{ms.split('(')[0].strip()}"); plt.close(fig_fim)
                st.download_button("⬇️ Feature Importance Excel", to_excel(fi_df),
                                   "feature_importance.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
            else:
                st.info("Feature Importance only available for RF and LDA.")

        # S.4 PCA-Raum ───────────────────────────────────────────────────────
        elif _page == SUP_PAGES[3]:
            st.title("🗺️ S.4 · PCA Space (Train + Test + Unknowns)")
            mz = st.session_state.mz_state
            if mz is None: st.warning("⚠️ Complete step S.1 first."); st.stop()
            ms = st.selectbox("Modell", list(mz["pipes"].keys()))
            pipe = mz["pipes"][ms]; features = mz["features"]
            Xtr = mz["X_train"]; Xte = mz["X_test"]
            ytr = mz["y_train"].astype(str).values; yte = mz["y_test"].astype(str).values
            ypred = mz["preds_test"][ms]
            df_unk = st.session_state.mz_df_unknown; Xunk = None
            if df_unk is not None:
                mf = [f for f in features if f not in df_unk.columns]
                if mf: st.warning(f"Bauteile fehlen: {mf}")
                else: Xunk = mz_lod_log(df_unk,features,mz.get("lod_map"),
                                         mz.get("lod_rule","lod2"),mz.get("do_log10",True),mz.get("eps",1e-12))
            if st.button("🗺️ Project", type="primary"):
                Ztr, Zte, Zunk, ve = mz_pca_proj(pipe, Xtr, Xte, Xunk)
                mis = (ypred.astype(str) != yte.astype(str))
                title = f"PCA · {ms} [{cl}]"
                idx_tr = mz["X_train"].index; idx_te = mz["X_test"].index
                df_ref_tr = mz["df_train"].loc[idx_tr, [HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)
                df_ref_te = mz["df_train"].loc[idx_te, [HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)
                Z_all = np.vstack([Ztr, Zte])
                sets_all = ["Train"]*len(Ztr) + ["Test"]*len(Zte)
                df_ref_all = pd.concat([df_ref_tr, df_ref_te], ignore_index=True)
                dfp = pd.concat([pd.DataFrame({"PC1":Z_all[:,0],"PC2":Z_all[:,1],"Set":sets_all}),
                                  df_ref_all], axis=1)
                st.subheader("🖱️ Interactive — Train + Test + Unknowns")
                fig_px = px.scatter(dfp, x="PC1", y="PC2", color=HERKUNFT_COL, symbol="Set",
                                    color_discrete_map=pal_color_map(p),
                                    title=title, template="plotly_white",
                                    labels={"PC1":f"PC1({ve[0]*100:.1f}%)","PC2":f"PC2({ve[1]*100:.1f}%)"})
                if mis.any():
                    fig_px.add_scatter(x=Zte[mis,0], y=Zte[mis,1], mode="markers",
                                       name="Misclassified",
                                       marker=dict(symbol="x",size=14,color="red",line_width=2))
                if Zunk is not None:
                    nc2 = st.session_state.mz_unk_name_col
                    un = (df_unk[nc2].astype(str).tolist() if nc2 and nc2 in df_unk.columns
                          else [f"B{i+1}" for i in range(len(df_unk))])
                    fig_px.add_scatter(x=Zunk[:,0], y=Zunk[:,1], mode="markers+text",
                                       name="Bauteile", text=un, textposition="top center",
                                       marker=dict(symbol="star",size=14,color="gold",line_color="black",line_width=1))
                fig_px.update_traces(marker_size=10, marker_opacity=0.88)
                st.plotly_chart(fig_px, use_container_width=True)
                st.divider(); st.subheader("🖨️ Publication — Train + Test + Unknowns")
                extras = {}
                if mis.any(): extras["Misclassified"] = (Zte[mis],"red","x",90)
                if Zunk is not None: extras["Bauteile"] = (Zunk,"gold","*",130)
                # Show TEST points too in pub plot
                fig_m_pca = pca_mpl(Ztr, df_ref_tr, p, cl, title, ve)
                ax_pub = fig_m_pca.axes[0]
                done_te = set()
                for i, (_, row_te) in enumerate(df_ref_te.iterrows()):
                    h_te = str(row_te[HERKUNFT_COL]); col_te, mkr_te = p.get(h_te,("#aaa","o"))
                    lbl_te = f"Test:{h_te}" if h_te not in done_te else "_nl_"; done_te.add(h_te)
                    ax_pub.scatter(Zte[i,0], Zte[i,1], color=col_te, marker=mkr_te,
                                   s=40, alpha=0.65, label=lbl_te, lw=0.3,
                                   edgecolors="black", zorder=4)
                if mis.any():
                    ax_pub.scatter(Zte[mis,0], Zte[mis,1], marker="x", s=90,
                                   color="red", label="Misclassified", lw=2, zorder=6)
                if Zunk is not None:
                    ax_pub.scatter(Zunk[:,0], Zunk[:,1], marker="*", s=130,
                                   color="gold", label="Bauteile", edgecolors="black", lw=0.8, zorder=7)
                ax_pub.legend(bbox_to_anchor=(1.02,1), loc="upper left", fontsize=6.5, framealpha=0.85)
                fig_m_pca.tight_layout()
                st.pyplot(fig_m_pca); pub_dl(fig_m_pca, f"pca_mz_{ms.split('(')[0].strip()}"); plt.close(fig_m_pca)
                c1m,c2m,c3m = st.columns(3)
                c1m.metric("PC1",f"{ve[0]*100:.1f}%"); c2m.metric("PC2",f"{ve[1]*100:.1f}%"); c3m.metric("Misclassified",int(mis.sum()))

        # S.5 Outlier & Gate ─────────────────────────────────────────────────
        elif _page == SUP_PAGES[4]:
            st.title("🚪 S.5 · Outlier Gate: Q-Residual + Hotelling T²")
            mz = st.session_state.mz_state; df_unk = st.session_state.mz_df_unknown
            if mz is None: st.warning("⚠️ Complete step S.1 first."); st.stop()
            if df_unk is None: st.warning("⚠️ Bauteile in S.1 hochladen."); st.stop()
            c1,c2,c3,c4 = st.columns(4)
            with c1: ms = st.selectbox("Modell",list(mz["pipes"].keys()))
            with c2: pt = st.slider("P(thr)",0.50,0.99,0.80,0.01)
            with c3: qq = st.slider("Q-Quantil",0.80,0.99,0.95,0.01)
            with c4: qp = st.slider("PCA Komp.",2,10,2)
            tk = st.slider("Top-k Klassen",2,5,3)
            st.info("**Gate Logic:** A sample is only assigned when **all 3** conditions are met:\n"
                    "1. Q-Residual ≤ threshold (within model space) \n"
                    "2. Hotelling T² ≤ threshold (not too far from center)\n"
                    "3. Class probability ≥ P(thr)")
            if st.button("🚪 Compute Gate", type="primary"):
                try:
                    tbl, s = mz_hybrid(mz, ms, df_unk, st.session_state.mz_unk_name_col, pt, qq, qp, tk)
                    cc1,cc2,cc3,cc4 = st.columns(4)
                    cc1.metric("Overall",s["n"]); cc2.metric("✅ Assigned",s["assigned"])
                    cc3.metric("❌ UNASSIGNED",s["unassigned"]); cc4.metric("Q-Thr",f"{s['q_threshold']:.3g}")
                    c5,c6 = st.columns(2)
                    c5.metric("T²-Thr",f"{s['T2_threshold']:.3g}"); c6.metric("Pass Q+T²",s["pass_Q"])

                    # Highlight UNASSIGNED reason
                    st.subheader("Results Table with Rejection Reason")
                    def highlight_gate(row):
                        if row["Assigned"]: return ["background-color:#d4edda"]*len(row)
                        return ["background-color:#f8d7da"]*len(row)
                    st.dataframe(tbl.style.apply(highlight_gate,axis=1), use_container_width=True)

                    # Pie chart
                    fig_pie = px.pie(values=[s["assigned"],s["unassigned"]],
                                     names=["Assigned","UNASSIGNED"],
                                     color_discrete_sequence=["#2ecc71","#e74c3c"],
                                     title="Assignment Rate")
                    st.plotly_chart(fig_pie, use_container_width=True)
                    # Gate breakdown
                    gate_summary = pd.DataFrame({
                        "Condition": ["Pass Q-Residual","Pass Hotelling T²","Pass Prob"],
                        "Anzahl":    [s["pass_Q"], s["pass_T2"], s["pass_Prob"]],
                        "Rate":      [f"{s['pass_Q']/s['n']*100:.1f}%",
                                      f"{s['pass_T2']/s['n']*100:.1f}%",
                                      f"{s['pass_Prob']/s['n']*100:.1f}%"]})
                    st.dataframe(gate_summary, use_container_width=True)
                    st.session_state.mz_state["hybrid_table"] = tbl
                    st.download_button("⬇️ Gate Excel", to_excel(tbl), "hybrid_gate.xlsx",
                                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                except Exception as e: st.error(f"Error: {e}")

        # S.6 Schwellwert ─────────────────────────────────────────────────────
        elif _page == SUP_PAGES[5]:
            st.title("📉 S.6 · Threshold Analysis")
            mz = st.session_state.mz_state
            if mz is None: st.warning("⚠️ Complete step S.1 first."); st.stop()
            c1,c2,c3 = st.columns(3)
            with c1: ms = st.selectbox("Modell",list(mz["pipes"].keys()))
            with c2:
                pm2 = st.slider("P min",0.10,0.90,0.50,0.05)
                px3 = st.slider("P max",0.50,0.99,0.95,0.01)
                st3 = st.slider("Steps",5,30,12)
            with c3:
                qq2 = st.slider("Q-Quantil",0.80,0.99,0.95,0.01,key="s6q")
                qp2 = st.slider("PCA Komp.",2,10,2,key="s6p")
            if st.button("📉 Compute", type="primary"):
                dfc = mz_thr_curve(mz, ms, np.linspace(pm2,px3,st3), qq2, qp2)
                f1i = px.line(dfc, x="coverage", y="acc_assigned", markers=True,
                               template="plotly_white", title="Coverage vs Accuracy")
                f2i = px.line(dfc, x="prob_threshold", y="coverage", markers=True,
                               template="plotly_white", title="Schwellwert vs Coverage")
                st.plotly_chart(f1i, use_container_width=True)
                st.plotly_chart(f2i, use_container_width=True)
                fig_thr, axes_thr = plt.subplots(1,2,figsize=(12,4.5))
                axes_thr[0].plot(dfc["coverage"],dfc["acc_assigned"],"o-",color=TAB10[0],lw=2,ms=6)
                axes_thr[0].set_xlabel("Coverage"); axes_thr[0].set_ylabel("Accuracy")
                axes_thr[0].set_title("Coverage vs Accuracy",fontweight="bold"); axes_thr[0].grid(True,alpha=0.2)
                axes_thr[1].plot(dfc["prob_threshold"],dfc["coverage"],"s-",color=TAB10[1],lw=2,ms=6)
                axes_thr[1].set_xlabel("Prob-Schwellwert"); axes_thr[1].set_ylabel("Coverage")
                axes_thr[1].set_title("Schwellwert vs Coverage",fontweight="bold"); axes_thr[1].grid(True,alpha=0.2)
                fig_thr.tight_layout(); st.pyplot(fig_thr)
                pub_dl(fig_thr, f"thr_{ms.split('(')[0].strip()}"); plt.close(fig_thr)
                st.dataframe(dfc.round(4), use_container_width=True)
                if st.session_state.mz_state: st.session_state.mz_state["threshold_curve"] = dfc

        # S.7 Wiederholung + Permutation ─────────────────────────────────────
        elif _page == SUP_PAGES[6]:
            st.title("🔁 S.7 · Repeated Validation + Permutation Test")
            mz = st.session_state.mz_state
            if mz is None or df_num is None: st.warning("⚠️ Complete step S.1 first."); st.stop()
            tab_rep, tab_perm = st.tabs(["🔁 Repeated Validation","🎲 Permutation Test"])

            with tab_rep:
                c1,c2 = st.columns(2)
                with c1: reps = st.slider("Repeats",5,200,30,5)
                with c2: rseed = st.number_input("Base Seed",value=int(mz.get("seed",42)),step=1)
                st.caption("Verwendet dieselbe Pipeline wie S.1 (gleiche Modelle, gleiches Preprocessing)")
                if st.button("🔁 Start", type="primary"):
                    with st.spinner(f"{reps}×4 …"):
                        try:
                            dr, ag = mz_repeated(mz, reps, rseed)
                            st.session_state.mz_state["repeated_runs"] = dr
                            st.session_state.mz_state["repeated_summary"] = ag
                            st.dataframe(ag.style.background_gradient(
                                subset=["macroF1_mean","accuracy_mean"],cmap="RdYlGn"),
                                use_container_width=True)
                            fig_rv = go.Figure()
                            for _, row in ag.iterrows():
                                fig_rv.add_bar(name=row["model"],x=[row["model"]],
                                               y=[row["macroF1_mean"]],
                                               error_y=dict(type="data",array=[row["macroF1_std"]],visible=True))
                            fig_rv.update_layout(template="plotly_white",title="MacroF1 mean±std",showlegend=False)
                            st.plotly_chart(fig_rv, use_container_width=True)
                            fig_rv2, ax_rv = plt.subplots(figsize=(7,4.5)); x_rv = np.arange(len(ag))
                            ax_rv.bar(x_rv, ag["macroF1_mean"], yerr=ag["macroF1_std"], capsize=5,
                                     color=[TAB10[i%10] for i in range(len(ag))], alpha=0.85, edgecolor="white")
                            ax_rv.set_xticks(x_rv); ax_rv.set_xticklabels(ag["model"],rotation=15,ha="right")
                            ax_rv.set_ylabel("MacroF1 (mean±std)")
                            ax_rv.set_title(f"Wiederholte Validierung (n={reps})",fontweight="bold")
                            ax_rv.set_ylim(0,min(1.05,ag["macroF1_mean"].max()*1.3))
                            ax_rv.grid(True,axis="y",alpha=0.2); fig_rv2.tight_layout()
                            st.pyplot(fig_rv2); pub_dl(fig_rv2,f"rep_n{reps}"); plt.close(fig_rv2)
                            st.download_button("⬇️ Excel",multi_excel({"Summary":ag,"All_Runs":dr}),
                                               "rep_val.xlsx","application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                        except Exception as e: st.error(f"Error: {e}")
                elif "repeated_summary" in (mz or {}):
                    st.dataframe(mz["repeated_summary"], use_container_width=True)

            with tab_perm:
                st.info("**Permutationstest:** Labels are randomly shuffled N times and the model is "
                        "retrained/evaluated. If the real model is significantly better than "
                        "chance, the classification is statistically robust.\n\n"
                        "p-Wert < 0.05 → Modell ist besser als Zufall")
                c1p,c2p,c3p = st.columns(3)
                with c1p: ms_perm = st.selectbox("Modell",list(mz["pipes"].keys()),key="perm_ms")
                with c2p: n_perm = st.slider("Permutations",49,499,99,50)
                with c3p: cv_perm = st.slider("CV Folds",3,10,5,key="perm_cv")
                if st.button("🎲 Run Permutation Test", type="primary"):
                    with st.spinner(f"{n_perm} Permutationen …"):
                        try:
                            score_real, perm_scores, pvalue = mz_permutation(mz, ms_perm, n_perm, cv_perm)
                            cp1,cp2,cp3 = st.columns(3)
                            cp1.metric("Echte Bal.Acc", f"{score_real:.4f}")
                            cp2.metric("Perm. Mittel",  f"{perm_scores.mean():.4f}")
                            cp3.metric("p-value", f"{pvalue:.4f}",
                                       delta="✅ significant" if pvalue < 0.05 else "❌ nicht signifikant")
                            if pvalue < 0.05:
                                st.success(f"✅ p={pvalue:.4f} — das Modell ist statistisch signifikant besser als Zufall!")
                            else:
                                st.warning(f"⚠️ p={pvalue:.4f} — kein signifikanter Unterschied zu Zufall (p ≥ 0.05)")
                            fig_perm, ax_perm = plt.subplots(figsize=(7,4))
                            ax_perm.hist(perm_scores, bins=25, color=TAB10[1], alpha=0.7, label="Permutiert")
                            ax_perm.axvline(score_real, color="red", lw=2, ls="-", label=f"Echt: {score_real:.4f}")
                            ax_perm.axvline(np.percentile(perm_scores,95), color="orange", lw=1.5,
                                            ls="--", label="95. Perzentil")
                            ax_perm.set_xlabel("Balanced Accuracy"); ax_perm.set_ylabel("Frequency")
                            ax_perm.set_title(f"Permutationstest — {ms_perm}\np={pvalue:.4f}", fontweight="bold")
                            ax_perm.legend(fontsize=9); ax_perm.grid(True, alpha=0.2)
                            fig_perm.tight_layout(); st.pyplot(fig_perm)
                            pub_dl(fig_perm, f"permtest_{ms_perm.split('(')[0].strip()}"); plt.close(fig_perm)
                        except Exception as e: st.error(f"Error: {e}")

        # S.8 Export ─────────────────────────────────────────────────────────
        elif _page == SUP_PAGES[7]:
            st.title("📤 S.8 · Complete Export")
            mz = st.session_state.mz_state
            if mz is None: st.warning("⚠️ Complete step S.1 first."); st.stop()
            st.subheader("Available Export Sheets")
            sheets = {}
            sheets["CV_Ergebnisse"]   = mz["cv_results"]
            sheets["Split_Ergebnisse"]= mz["split_results"]
            features_df = pd.DataFrame({"Feature":mz["features"]})
            sheets["Features"] = features_df
            if "hybrid_table" in mz and mz["hybrid_table"] is not None:
                sheets["Hybrid_Gate"] = mz["hybrid_table"]
            if "threshold_curve" in mz and mz["threshold_curve"] is not None:
                sheets["Schwellwert"] = mz["threshold_curve"]
            if "repeated_summary" in mz and mz["repeated_summary"] is not None:
                sheets["Wiederholt_Summary"]  = mz["repeated_summary"]
                sheets["Wiederholt_Alle_Runs"] = mz["repeated_runs"]
            if "plsda" in mz and mz["plsda"] is not None:
                sheets["PLSDA_VIP"] = mz["plsda"]["vip_df"]
            for ms_i in mz["pipes"].keys():
                fi_df = mz_feature_importance(mz, ms_i)
                if fi_df is not None:
                    sheets[f"FeatImp_{ms_i[:15]}"] = fi_df
                lbls, cm_i = mz["confmats"][ms_i]
                cm_df = pd.DataFrame(cm_i, index=[f"T:{l}" for l in lbls],
                                     columns=[f"P:{l}" for l in lbls])
                cm_df.index.name = "True"
                sheets[f"CM_{ms_i[:15]}"] = cm_df.reset_index()
            for nm, df_s in sheets.items():
                st.markdown(f"- **{nm}** ({len(df_s)} Zeilen)")
            st.divider()
            excel_bytes = multi_excel(sheets)
            st.download_button("📥 Complete Excel Export",
                               excel_bytes, "chemotrace_export.xlsx",
                               "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                               use_container_width=True)
            if df_num is not None:
                st.download_button("📥 Raw Data + Classes Excel",
                                   to_excel(df_num), "raw_data_classes.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                   use_container_width=True)

    # ══════ UNSUPERVISED ═════════════════════════════════════════════════════
    elif "Unsupervised" in _mz_mode:
        if st.session_state.df_proc is None:
            st.error("❌ Preprocessing (1.3) abschließen."); st.stop()
        d = st.session_state.df_proc; feats = st.session_state.feature_cols
        X = SimpleImputer(strategy="median").fit_transform(d[feats].astype(float).values)

        # U.1 PCA Biplot ─────────────────────────────────────────────────────
        if _page == UNS_PAGES[0]:
            st.title("🌀 U.1 · PCA Biplot (Unsupervised)")
            with st.expander("ℹ️ Methodik: PCA Biplot", expanded=False):
                st.markdown("""
A **PCA biplot** overlays the scores plot (sample positions) with loading vectors (variable
contributions). Arrow direction indicates the influence axis of each element; arrow length
reflects the magnitude of its contribution. This combined view enables simultaneous
interpretation of sample groupings and the chemical variables driving the separation — a key
tool in exploratory chemometric provenance analysis.
                """)
            st.info(f"Class colors for **orientation** — Level: **{cl}** · "
                    f"Symbol: **{'Sample Origin' if st.session_state.symbol_detail else 'uniform'}**")
            c1,c2 = st.columns(2)
            with c1: top_n = st.slider("Loading Arrows",3,20,8); scale2 = st.slider("Arrow Scaling",0.5,8.0,3.0,0.5)
            with c2: top_ld2 = st.slider("Top Loadings Balken",5,25,12)
            if st.button("🌀 Compute PCA", type="primary"):
                Z,ve,ld_df,_ = run_pca(X, feats)
                ld2 = ld_df[["PC1","PC2"]].copy(); ld2["Feature"] = ld_df.index
                ld2["max_abs"] = ld2[["PC1","PC2"]].abs().max(axis=1)
                sc_df = pd.concat([pd.DataFrame({"PC1":Z[:,0],"PC2":Z[:,1]}),
                                    d[[HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)], axis=1)
                st.subheader("🖱️ Interactive")
                top_f = ld_df["PC1"].abs().add(ld_df["PC2"].abs()).nlargest(top_n).index
                fig_upca = px.scatter(sc_df, x="PC1", y="PC2", color=HERKUNFT_COL, symbol=HERKUNFT_COL,
                                      color_discrete_map=pal_color_map(p), symbol_map=pal_symbol_map(p),
                                      hover_data=CLASS_LEVELS,
                                      title=f"PCA Unsupervised [{cl}]", template="plotly_white",
                                      labels={"PC1":f"PC1({ve[0]*100:.1f}%)","PC2":f"PC2({ve[1]*100:.1f}%)"})
                for feat in top_f:
                    lx = ld_df.loc[feat,"PC1"]*scale2; ly = ld_df.loc[feat,"PC2"]*scale2
                    fig_upca.add_annotation(x=lx,y=ly,ax=0,ay=0,xref="x",yref="y",axref="x",ayref="y",
                                            showarrow=True,arrowhead=2,arrowwidth=1.5,arrowcolor="#555")
                    fig_upca.add_annotation(x=lx*1.1,y=ly*1.1,text=feat,showarrow=False,
                                            font=dict(size=10,color="#333"))
                fig_upca.update_traces(marker_size=10,marker_opacity=0.88)
                st.plotly_chart(fig_upca, use_container_width=True)
                st.divider(); st.subheader("🖨️ Biplot")
                fig_ubp = biplot_mpl(Z,ld_df,d.reset_index(drop=True),p,cl,f"PCA Biplot [{cl}]",ve,top_n,scale2)
                st.pyplot(fig_ubp); pub_dl(fig_ubp,"pca_biplot_unsup"); plt.close(fig_ubp)
                st.subheader("🖨️ Loadings")
                fig_uld = loadings_mpl(ld2,top_ld2,"Top Loadings")
                st.pyplot(fig_uld); pub_dl(fig_uld,"pca_loadings_unsup"); plt.close(fig_uld)
                c1m,c2m = st.columns(2); c1m.metric("PC1",f"{ve[0]*100:.1f}%"); c2m.metric("PC2",f"{ve[1]*100:.1f}%")

        # U.2 k-Means ────────────────────────────────────────────────────────
        elif _page == UNS_PAGES[1]:
            st.title("🔵 U.2 · k-Means Clustering")
            with st.expander("ℹ️ Methodik: k-Means Clustering", expanded=False):
                st.markdown("""
**k-Means** is an unsupervised partitioning algorithm that assigns *n* samples to *k* clusters
by minimizing within-cluster sum of squares (WCSS). The **Elbow method** identifies the optimal *k*
by plotting WCSS against cluster count. The **Silhouette score** (range −1 to +1) measures how
similar a sample is to its own cluster versus neighboring clusters — values above 0.5 indicate
good separation.
                """)
            st.info(f"Colors = known classes [{cl}] for orientation. Cluster number as text.")
            c1,c2,c3 = st.columns(3)
            with c1: k = st.slider("k",2,min(12,len(d)-1),4)
            with c2: km_seed = st.number_input("Seed",42,step=1,key="kms2")
            with c3: elbow = st.checkbox("Elbow + Silhouette",True)
            if st.button("🔵 Compute", type="primary"):
                km = KMeans(n_clusters=k, random_state=int(km_seed), n_init="auto")
                km_labels = km.fit_predict(X)
                sil_score = silhouette_score(X, km_labels) if len(set(km_labels)) > 1 else np.nan
                st.metric("Silhouette Score", f"{sil_score:.4f}",
                          delta="gut (>0.5)" if sil_score > 0.5 else ("ok (0.25-0.5)" if sil_score > 0.25 else "schwach (<0.25)"))
                if elbow:
                    ks = range(2, min(13,len(d))); inis = []; sils = []
                    for ki in ks:
                        km2 = KMeans(n_clusters=ki,random_state=int(km_seed),n_init="auto")
                        lbl2 = km2.fit_predict(X); inis.append(km2.inertia_)
                        sils.append(silhouette_score(X,lbl2) if len(set(lbl2))>1 else 0)
                    fig_el, (ax_el1,ax_el2) = plt.subplots(1,2,figsize=(10,3.5))
                    ax_el1.plot(list(ks),inis,"o-",color=TAB10[0],lw=2); ax_el1.axvline(k,color="red",ls="--",alpha=0.6,label=f"k={k}")
                    ax_el1.set_xlabel("k"); ax_el1.set_ylabel("WCSS"); ax_el1.set_title("Elbow",fontweight="bold"); ax_el1.legend(); ax_el1.grid(True,alpha=0.2)
                    ax_el2.plot(list(ks),sils,"s-",color=TAB10[1],lw=2); ax_el2.axvline(k,color="red",ls="--",alpha=0.6,label=f"k={k}")
                    ax_el2.set_xlabel("k"); ax_el2.set_ylabel("Silhouette"); ax_el2.set_title("Silhouette Score",fontweight="bold"); ax_el2.legend(); ax_el2.grid(True,alpha=0.2)
                    fig_el.tight_layout(); st.pyplot(fig_el); plt.close(fig_el)
                Z_pca,ve_pca,_,_ = run_pca(X,feats)
                sc_df = pd.concat([pd.DataFrame({"PC1":Z_pca[:,0],"PC2":Z_pca[:,1],"Cluster":km_labels.astype(str)}),
                                    d[[HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)], axis=1)
                st.subheader("🖱️ Interactive")
                fig_km_px = px.scatter(sc_df, x="PC1", y="PC2", color=HERKUNFT_COL, symbol="Cluster",
                                        color_discrete_map=pal_color_map(p), symbol_map=pal_symbol_map(p),
                                        hover_data=CLASS_LEVELS+["Cluster"],
                                        title=f"k-Means k={k} [{cl}] · Silhouette={sil_score:.3f}",
                                        template="plotly_white")
                fig_km_px.update_traces(marker_size=11,marker_opacity=0.88)
                st.plotly_chart(fig_km_px, use_container_width=True)
                st.divider(); st.subheader("🖨️ Publication")
                fig_km2 = unsup_scatter_mpl(Z_pca,d.reset_index(drop=True),p,cl,
                                             f"k-Means k={k} [{cl}]\nSilhouette={sil_score:.3f}",
                                             f"PC1({ve_pca[0]*100:.1f}%)",f"PC2({ve_pca[1]*100:.1f}%)",
                                             cluster_labels=km_labels)
                st.pyplot(fig_km2); pub_dl(fig_km2,f"kmeans_k{k}"); plt.close(fig_km2)
                df_cl = d[[HERKUNFT_COL]+CLASS_LEVELS].copy(); df_cl["Cluster"] = km_labels
                st.subheader("Cluster Composition")
                st.dataframe(df_cl.groupby(["Cluster",cl])[HERKUNFT_COL].count().reset_index()
                              .rename(columns={HERKUNFT_COL:"Anzahl"}), use_container_width=True)
                st.download_button("⬇️ Cluster Excel",to_excel(df_cl),"kmeans.xlsx",
                                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # U.3 Dendrogram ─────────────────────────────────────────────────────
        elif _page == UNS_PAGES[2]:
            st.title("🌲 U.3 · Hierarchical Clustering")
            st.info(f"Leaf colors = known classes [{cl}].")
            c1,c2 = st.columns(2)
            with c1: method = st.selectbox("Linkage",["ward","complete","average","single"])
            with c2: metric = st.selectbox("Distanz",["euclidean","cosine","cityblock"])
            if method == "ward" and metric != "euclidean":
                st.warning("Ward → nur euclidean."); metric = "euclidean"
            if st.button("🌲 Dendrogram", type="primary"):
                Zl = linkage(X, method=method, metric=metric)
                labels_dn = d[HERKUNFT_COL].astype(str).tolist()
                st.subheader("🖨️ Dendrogram")
                fig_dn = dendro_mpl(Zl, labels_dn, p, f"Hierarchisch ({method}/{metric}) [{cl}]")
                st.pyplot(fig_dn); pub_dl(fig_dn,f"dendro_{method}"); plt.close(fig_dn)
                ncut = st.slider("Cluster nach Cut",2,min(12,len(d)-1),4)
                ac = AgglomerativeClustering(n_clusters=ncut, linkage=method,
                                             metric=metric if method!="ward" else "euclidean")
                cut_l = ac.fit_predict(X)
                sil_c = silhouette_score(X,cut_l) if len(set(cut_l))>1 else np.nan
                st.metric("Silhouette Score",f"{sil_c:.4f}")
                Z_pca2,ve_pca2,_,_ = run_pca(X,feats)
                sc_df2 = pd.concat([pd.DataFrame({"PC1":Z_pca2[:,0],"PC2":Z_pca2[:,1],"Cluster":cut_l.astype(str)}),
                                     d[[HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)], axis=1)
                fig_dn_sc = px.scatter(sc_df2,x="PC1",y="PC2",color=HERKUNFT_COL,symbol="Cluster",
                                        color_discrete_map=pal_color_map(p),
                                        title=f"Hierarchisch {ncut} Cluster [{cl}] · Silhouette={sil_c:.3f}",
                                        template="plotly_white")
                fig_dn_sc.update_traces(marker_size=11,marker_opacity=0.88)
                st.plotly_chart(fig_dn_sc, use_container_width=True)

        # U.4 UMAP ───────────────────────────────────────────────────────────
        elif _page == UNS_PAGES[3]:
            st.title("🗺️ U.4 · UMAP")
            st.warning("⚠️ **UMAP is an exploratory tool.** The axes are not quantitatively "
                       "interpretierbar (anders als PCA). Abstände und Winkel im UMAP-Plot dürfen NICHT "
                       "wie PCA-Loadings gelesen werden. Für quantitative Aussagen → PCA verwenden.")
            if not UMAP_AVAILABLE:
                st.error("UMAP not installed."); st.code("pip install umap-learn"); st.stop()
            st.info(f"Colors = classes [{cl}] for orientation. UMAP has no access to labels.")
            c1,c2,c3 = st.columns(3)
            with c1: nn = st.slider("n_neighbors",5,50,15)
            with c2: md = st.slider("min_dist",0.01,0.99,0.10,0.01)
            with c3: us = st.number_input("Seed",42,step=1,key="us2")
            if st.button("🗺️ Compute UMAP", type="primary"):
                with st.spinner("UMAP …"):
                    reducer = umap_lib.UMAP(n_neighbors=nn,min_dist=md,random_state=int(us))
                    Zu = reducer.fit_transform(X)
                sc_df3 = pd.concat([pd.DataFrame({"UMAP1":Zu[:,0],"UMAP2":Zu[:,1]}),
                                     d[[HERKUNFT_COL]+CLASS_LEVELS].reset_index(drop=True)], axis=1)
                st.subheader("🖱️ Interactive")
                fig_um = px.scatter(sc_df3, x="UMAP1", y="UMAP2", color=HERKUNFT_COL, symbol=HERKUNFT_COL,
                                     color_discrete_map=pal_color_map(p), symbol_map=pal_symbol_map(p),
                                     hover_data=CLASS_LEVELS,
                                     title=f"UMAP [{cl}] n_nbr={nn} dist={md}",
                                     template="plotly_white")
                fig_um.update_traces(marker_size=11,marker_opacity=0.88)
                st.plotly_chart(fig_um, use_container_width=True)
                st.divider(); st.subheader("🖨️ Publication")
                fig_um2 = unsup_scatter_mpl(Zu,d.reset_index(drop=True),p,cl,
                                             f"UMAP [{cl}]\n(explorativ — Achsen nicht quantitativ interpretierbar)",
                                             "UMAP 1","UMAP 2")
                st.pyplot(fig_um2); pub_dl(fig_um2,f"umap_n{nn}"); plt.close(fig_um2)
